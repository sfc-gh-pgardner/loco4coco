#!/usr/bin/env python3
"""Loco4CoCo arcade companion - local server.

The visitor drives a penguin round an arctic map. At each location the browser
posts their choices here; this server prompts the REAL Cortex Code via
`cortex exec` and streams what CoCo is doing back to the screen. Nothing in the
browser talks to Snowflake or to an MCP directly.

    browser --(HTTP)--> server --(cortex exec)--> CoCo --> Snowflake / Gmail MCP

Two exec profiles, deliberately different:

  conversational turns : --no-mcp   (no tools needed, and skipping MCP startup
                                     is the single biggest latency win)
  the postbox send     : --bypass   (MCP on; without --bypass the Gmail tool
                                     call is AUTO-REJECTED in exec mode)

Delivery is Route A behind a transport interface. See config.delivery.

  python3 server.py                 # use config.json
  python3 server.py --port 5000     # override
"""

import argparse
import base64
import binascii
import hashlib
import html
import json
import os
import posixpath
import re
import subprocess
import tempfile
import threading
import time
import uuid
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

# The booth's closed lists, with their own three-layer fallback (Snowflake, the
# committed bundle, then this repo's markdown). Kept in a separate module because
# the ops pre-flight and the loader both need it without starting a web server.
import context as bctx
import agent_pool
from urllib.parse import parse_qs, urlparse

HERE = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(HERE, "config.json")
STATE_PATH = os.path.join(HERE, "state.json")
PLUGIN_ROOT = os.path.dirname(HERE)
GUIDES_PATH = os.path.join(PLUGIN_ROOT, "skills", "loco4coco",
                           "references", "guides-index.md")
FEATURES_PATH = os.path.join(PLUGIN_ROOT, "skills", "loco4coco",
                             "references", "feature-docs.md")
MARKET_PATH = os.path.join(PLUGIN_ROOT, "skills", "loco4coco",
                           "references", "marketplace-index.md")
LISTING_URL = "https://app.snowflake.com/marketplace/listing/"

_lock = threading.Lock()

MIME = {
    ".html": "text/html; charset=utf-8",
    ".js": "text/javascript; charset=utf-8",
    ".css": "text/css; charset=utf-8",
    ".json": "application/json; charset=utf-8",
    ".png": "image/png",
    ".svg": "image/svg+xml",
}

BLANK_STATE = {
    "stage": "attract",
    "session_id": "",
    "location": "",
    "thinking": False,
    "transport": "",
    "reasoning": [],
    "reply": "",
    "narrative": "",
    "turns": [],
    "visitor": {"first_name": "", "company": "", "email": "", "industry": "",
                "problem": "", "company_country": "", "residency": ""},
    "held": [],
    "platforms": [],
    "joined": [],
    "joined_listings": [],
    "poc": {},
    "unlocked": ["library"],
    "blueprint_url": "",
    "card_url": "",
    "draft_created": False,
    "queued": False,
    "email_sent": False,
    "logged": False,
    "ask_used": False,
    "log_error": "",
    "coco_seconds": 0,
    "input_tokens": 0,
    "output_tokens": 0,
    "job_id": None,
    "started_at": None,
    "updated_at": None,
}

# Coarse industry inference. Instant and free, which matters against a
# five-minute clock; the visitor always gets to correct it with one click.
INDUSTRY_HINTS = {
    "healthcare": ["nhs", "hospital", "trust", "health", "clinic", "care",
                   "medical", "pharma", "biotech", "gp ", "surgery", "patient"],
    "financial": ["bank", "capital", "invest", "insur", "asset", "fund",
                  "lloyds", "barclays", "hsbc", "natwest", "santander",
                  "financ", "fintech", "pension", "building society", "credit"],
    "retail": ["retail", "tesco", "sainsbury", "asda", "morrison", "aldi",
               "lidl", "marks", "boots", "shop", "store", "consumer",
               "unilever", "nestle", "diageo", "brand", "grocer", "fashion"],
    "public": ["council", "borough", "government", "ministry", "department",
               "authority", "police", "dwp", "hmrc", "defra", "agency",
               "public", "county", "city of", "gov.uk", "school", "university"],
    "manufacturing": ["manufactur", "factory", "industri", "engineer",
                      "automotive", "aerospace", "rolls", "bae", "jaguar",
                      "steel", "chemical", "plant", "production"],
    "energy": ["energy", "power", "utilit", "water", "grid", "national grid",
               "shell", "bp ", "centrica", "octopus", "sse", "electric",
               "gas", "renewab", "nuclear", "thames"],
    "media": ["media", "broadcast", "bbc", "sky", "itv", "telecom", "vodafone",
              "bt ", "o2", "three", "ee ", "publish", "music", "studio",
              "entertain", "advertis", "agency", "game"],
}


def load_config():
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return json.load(f)


def read_state():
    with _lock:
        if not os.path.exists(STATE_PATH):
            return dict(BLANK_STATE)
        try:
            with open(STATE_PATH, encoding="utf-8") as f:
                data = json.load(f)
        except (json.JSONDecodeError, OSError):
            # A half-written file must never crash the booth display.
            return dict(BLANK_STATE)
    merged = dict(BLANK_STATE)
    merged.update(data)
    return merged


def purge_temp_artifacts():
    """Delete the visitor-named blueprint files this process writes to the OS
    temp dir (loco4coco-<name>-<ts>.html/.docx). They carry the visitor's first
    name, so on a shared booth they must not outlive the visit - called on reset
    so nothing about one visitor is left on disk for the next. Returns a count."""
    removed = 0
    tmp = tempfile.gettempdir()
    try:
        for n in os.listdir(tmp):
            if n.startswith("loco4coco-") and n.endswith((".html", ".docx")):
                try:
                    os.remove(os.path.join(tmp, n))
                    removed += 1
                except OSError:
                    pass
    except OSError:
        pass
    return removed


def write_state(patch, replace=False):
    """Shallow-merge a patch into state.json, atomically.

    `replace` writes the patch wholesale. Needed for reset: `visitor` and `poc`
    are merged sub-dicts, so passing an empty one would leave the previous
    visitor's POC in place - which is exactly how one visitor's name leaked
    into the next visitor's session.
    """
    with _lock:
        current = dict(BLANK_STATE)
        if os.path.exists(STATE_PATH) and not replace:
            try:
                with open(STATE_PATH, encoding="utf-8") as f:
                    current.update(json.load(f))
            except (json.JSONDecodeError, OSError):
                pass
        for k, v in patch.items():
            if k in ("visitor", "poc") and isinstance(v, dict) and not replace:
                sub = dict(current.get(k) or {})
                sub.update(v)
                current[k] = sub
            else:
                current[k] = v
        if current.get("stage") not in (None, "attract") and not current.get("started_at"):
            current["started_at"] = time.time()
        current["updated_at"] = time.time()
        tmp = STATE_PATH + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(current, f, indent=1)
        os.replace(tmp, STATE_PATH)
    return current


# ---------------------------------------------------------------- cost logging

def log_cost(cfg, kind, seconds, usage, ok, transport="exec", model=None):
    """Append-only spend trail with the real token counts exec reports. We do
    not cap the visitor, so this is the only way spend stays visible."""
    path = os.path.join(HERE, (cfg.get("coco") or {}).get("cost_log", "cost.jsonl"))
    u = usage or {}
    rec = {"at": time.time(), "kind": kind, "seconds": round(seconds, 2),
           "transport": transport, "model": model,
           "input_tokens": u.get("input_tokens"),
           "output_tokens": u.get("output_tokens"),
           "cache_read": u.get("cache_read_input_tokens"),
           "cache_write": u.get("cache_creation_input_tokens"),
           "ok": ok}
    try:
        with open(path, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec) + "\n")
    except OSError:
        pass


# ------------------------------------------------------------- the CoCo bridge

# Reasoning worth showing a visitor. Deliberately excludes assistant text,
# because that is the reply and it already has a home in the thought bubble.
def _readable(name):
    return {"sql_execute": "querying Snowflake", "web_search": "searching the web",
            "web_fetch": "reading a page", "read": "reading a file",
            "bash": "running a command", "python_repl": "working it out"}.get(
                name, name.replace("_", " "))


# Lines that break character or expose plumbing. A booth audience must never
# read "the user wants me to respond as CoCo the Snowflake penguin" - it tells
# them he was told to act like one - nor see internal tool names.
_TRAY_BLOCK = re.compile(
    r"(the user wants|the user is asking|the user said|i need to (keep|pick|respond)"
    r"|i should (keep|respond|reply)|no bullet points|plain english|short sentences"
    r"|max \d+ words|as coco|the coco|my instructions|the prompt|mcp__|create_draft"
    r"|tool search|toolset|google-workspace|google-calendar|gmail|skill\b"
    r"|verbatim|minified json|booth)", re.I)


def _sentences(text, limit=3):
    out = []
    for s in re.split(r"(?<=[.!?])\s+", (text or "").strip()):
        s = s.strip().lstrip("-*\u2022 ").strip()
        if len(s) < 12 or _TRAY_BLOCK.search(s):
            continue
        out.append(s)
        if len(out) >= limit:
            break
    return out


def run_exec(cfg, prompt, kind, job_id=None, use_mcp=False, timeout=None,
             effort=None, tools=None, max_turns=None, allowed=None):
    """Run `cortex exec` and stream its real reasoning into state.reasoning.

    Uses --format json, which emits NDJSON events including genuine `thinking`
    blocks, tool calls, and token usage. The tray therefore shows what CoCo is
    actually doing rather than a decorative spinner.
    """
    c = cfg.get("coco") or {}
    cmd = [c.get("binary", "cortex"), "exec", prompt, "--format", "json"]
    conn = (cfg.get("snowflake") or {}).get("connection_name")
    if conn:
        cmd += ["--connection", conn]
    if use_mcp:
        # Without --bypass, exec auto-rejects the Gmail tool call and the send
        # silently does nothing. Proven: "Tool denied: headless mode requires..."
        cmd.append("--bypass")
    else:
        cmd.append("--no-mcp")
        if tools:
            # Let CoCo actually USE its built-in tools. exec auto-rejects tool
            # calls without --bypass, so without this the tray only ever shows
            # thinking. MCP stays OFF: we want the built-ins, not the MCP
            # server startup cost.
            cmd.append("--bypass")
    if max_turns:
        # Safety ceiling on the agentic loop once tools are in play.
        cmd += ["--max-turns", str(max_turns)]
    if allowed:
        # Optional allow-list, e.g. ["Bash(gh *)"]. Unused by default; the
        # blueprint is clamped server-side anyway (off-list features dropped,
        # guide choice falls back), so a wandering model cannot leak through.
        cmd += ["--allowed", ",".join(allowed) if isinstance(allowed, list)
                else str(allowed)]
    # Each turn is a fresh, stateless exec: --no-history skips the session save
    # and keeps one visitor's turns from bleeding into the next.
    cmd.append("--no-history")
    # Booth-tunable model + effort. Global config.coco.effort is the default; a
    # per-turn `effort` override (from the location) wins. Measured: default
    # (null) behaves like high effort; "low" cuts ~8s off simple turns and ~17s
    # off the workshop, and the {minimal,low,medium} spread is only ~3-4s.
    model = c.get("model")
    if model:
        cmd += ["-m", str(model)]
    eff = effort or c.get("effort")
    if eff:
        cmd += ["--effort", str(eff)]

    started = time.time()
    lines, final, usage, ok = [], "", {}, False

    def push(line):
        lines.append(line)
        write_state({"reasoning": lines[-40:]})

    try:
        p = subprocess.Popen(cmd, stdout=subprocess.PIPE,
                             stderr=subprocess.DEVNULL, text=True,
                             bufsize=1, cwd=HERE)
        deadline = started + (timeout or c.get("timeout_seconds", 180))
        for raw in p.stdout:
            raw = raw.strip()
            if not raw:
                continue
            if job_id is not None:
                cur = read_state()
                if cur.get("job_id") not in (None, job_id):
                    p.kill()
                    return False, ""      # superseded by a newer message
            try:
                ev = json.loads(raw)
            except json.JSONDecodeError:
                continue
            t = ev.get("type")
            if t == "system" and ev.get("subtype") == "init":
                push("Processing your request...")
            elif t == "assistant":
                for b in (ev.get("message", {}).get("content") or []):
                    bt = b.get("type")
                    if bt == "thinking":
                        for s in _sentences(b.get("thinking")):
                            push(s)
                    elif bt == "tool_use":
                        push("\u2192 " + _readable(b.get("name") or "a tool"))
            elif t == "result":
                final = (ev.get("result") or "").strip()
                usage = ev.get("usage") or {}
                ok = not ev.get("is_error")
            if time.time() > deadline:
                p.kill()
                push("That is taking a while. I will go with what I have.")
                break
        p.wait(timeout=10)
        if p.returncode == 0 and final:
            ok = True
    except FileNotFoundError:
        return False, "CoCo CLI not found on PATH, so I cannot answer."
    except Exception as e:                                       # noqa: BLE001
        return False, f"Something went wrong reaching CoCo: {e}"

    log_cost(cfg, kind, time.time() - started, usage, ok, "exec", model)
    meta = {"seconds": round(time.time() - started, 1), "usage": usage,
            "ok": ok, "transport": "exec", "model": model}
    if not final:
        return False, "I lost my train of thought there. Try me again.", meta
    return ok, final, meta


# --------------------------------------------------- the fast path (no agent)

_conn = None
_conn_lock = threading.Lock()


def sf_conn(cfg):
    """One long-lived Snowflake connection, reused for Cortex COMPLETE calls and
    for row logging.

    Reusing a connection is the whole point: a fresh `snow sql` subprocess costs
    ~3.4s of CLI startup per statement, which is more than a COMPLETE call takes.
    """
    global _conn
    with _conn_lock:
        if _conn is not None:
            try:
                if not _conn.is_closed():
                    return _conn
            except Exception:                                    # noqa: BLE001
                pass
            _conn = None
        import snowflake.connector as sc                    # lazy: booth-only dep
        name = (cfg.get("snowflake") or {}).get("connection_name")
        _conn = sc.connect(**({"connection_name": name} if name else {}))
        return _conn


def run_complete(cfg, prompt, kind, job_id=None, model=None, timeout=None):
    """Call Cortex inference directly, in-process, with no agentic loop.

    This skips the ~22-25s fixed floor that `cortex exec` pays on every call
    (process spawn + connection + model attach), which is why the reflection
    turns land in ~1-2s instead of ~24-34s.

    There is genuinely no reasoning to stream here, so the tray is CLEARED
    rather than filled with a decorative spinner: the app only ever shows real
    thinking. On any failure this returns ok=False so the caller can fall back
    to exec (Claude was unavailable via COMPLETE on this region once already).
    """
    c = cfg.get("coco") or {}
    mdl = model or c.get("complete_model") or "mistral-large2"
    started = time.time()
    write_state({"reasoning": []})
    reply, err, usage = "", "", {}
    try:
        cur = sf_conn(cfg).cursor()
        try:
            # The three-argument form, not COMPLETE(model, prompt). The two-arg
            # version returns a bare string with no usage, which is why every
            # COMPLETE turn logged 0 input and 0 output tokens while the agentic
            # turn logged real ones - so "cost per visitor" was only ever
            # counting the Workshop. This form returns choices plus usage.
            cur.execute(
                "SELECT SNOWFLAKE.CORTEX.COMPLETE(%s, "
                "ARRAY_CONSTRUCT(OBJECT_CONSTRUCT('role','user',"
                "'content',%s)), OBJECT_CONSTRUCT()) ", (mdl, prompt))
            row = cur.fetchone()
            raw = row[0] if row else ""
            if isinstance(raw, str):
                raw = json.loads(raw) if raw.strip().startswith("{") else {}
            if isinstance(raw, dict):
                ch = (raw.get("choices") or [{}])[0]
                reply = str(ch.get("messages") or ch.get("message") or "").strip()
                u = raw.get("usage") or {}
                usage = {"input_tokens": int(u.get("prompt_tokens") or 0),
                         "output_tokens": int(u.get("completion_tokens") or 0)}
            else:
                reply = str(raw or "").strip()
        finally:
            cur.close()
    except Exception as e:                                       # noqa: BLE001
        err = str(e)[:200].replace("\n", " ")

    secs = time.time() - started
    ok = bool(reply) and not err
    log_cost(cfg, kind, secs, usage, ok, "complete", mdl)
    meta = {"seconds": round(secs, 1), "usage": usage, "ok": ok,
            "transport": "complete", "model": mdl}
    if err:
        meta["error"] = err
    if job_id is not None:
        cur_state = read_state()
        if cur_state.get("job_id") not in (None, job_id):
            return False, "", meta          # superseded by a newer message
    if not ok:
        return False, "", meta
    return True, reply, meta


def turn_budget(cfg, loc, timeout=None):
    """Wall-clock ceiling for one turn, in seconds.

    A ceiling is not a nicety. The Library stop has been measured at a 127.3s
    outlier against a 2.2s median, on a stop budgeted for a few seconds, inside a
    five-minute visit with a queue waiting. Past the ceiling the visitor is
    better served by the precomputed answer than by a better sentence.
    """
    coco = cfg.get("coco") or {}
    return int(timeout
               or loc.get("timeout")
               or coco.get("turn_timeout")
               or 60)


def run_turn(cfg, prompt, kind, loc=None, job_id=None, use_mcp=False,
             timeout=None):
    """Dispatch a turn, with four layers of fallback beneath it.

        1. warm agent   `cortex mcp serve`, held open between visitors. ~3.4s.
        2. cortex exec  a cold one-shot process. ~26s, of which ~18s is startup.
        3. COMPLETE     SNOWFLAKE.CORTEX.COMPLETE. Fast, non-agentic.
        4. precomputed  the archetype's own defaults, no model at all.

    Which layer leads depends on what the location asks for. A reflection turn
    wants COMPLETE and does not need judgement; the Workshop wants judgement and
    so leads with the warm agent. Every layer falls through, so a flat network, a
    suspended warehouse or a model outage degrades the visit rather than ending
    it - layer 4 needs nothing but this repo.

    Returns (ok, reply, meta). meta["transport"] records which layer answered,
    which is what makes the cost log worth reading afterwards.
    """
    loc = loc or {}
    budget = turn_budget(cfg, loc, timeout)
    coco = cfg.get("coco") or {}
    wants = (loc.get("transport") or "exec")
    started = time.time()
    meta = None

    def remaining():
        return max(4, budget - int(time.time() - started))

    if wants == "complete":
        ok, reply, meta = run_complete(cfg, prompt, kind, job_id=job_id,
                                       model=loc.get("complete_model"))
        if ok or meta.get("ok"):
            return ok, reply, meta
        if not meta.get("error"):
            return ok, reply, meta          # superseded, not a failure
        print(f"[loco4coco] COMPLETE failed for {kind}, falling back: "
              f"{meta['error']}")

    # Layer 1: the warm agent. Skipped when disabled in config, and skipped
    # quietly when the process will not start - nobody at the stand should ever
    # be able to tell which layer answered.
    elif coco.get("warm_agent", True):
        agent = agent_pool.get_agent(
            connection=coco_connection(), model=coco.get("model"),
            workdir=HERE, log=lambda m: print("[loco4coco] " + m))
        if agent.alive() or agent.start():
            ok, reply, meta = agent.ask(prompt, timeout=remaining())
            # log_cost lives inside run_exec and run_complete, so the warm path
            # has to record its own turn or the Workshop - the stop this whole
            # transport exists for - would vanish from the cost log entirely.
            log_cost(cfg, kind, meta.get("seconds") or 0, {}, ok, "warm",
                     coco.get("model"))
            if ok:
                return True, reply, meta
            print("[loco4coco] warm agent failed for %s (%s), falling back to "
                  "exec" % (kind, meta.get("error")))

    # Layer 2: a cold one-shot process. Needs ~20s just to start, so it is not
    # worth beginning if the budget cannot cover that.
    if remaining() > 20:
        ok, reply, meta = run_exec(
            cfg, prompt, kind, job_id=job_id, use_mcp=use_mcp,
            timeout=remaining(), effort=loc.get("effort"),
            tools=loc.get("tools"), max_turns=loc.get("max_turns"))
        if ok:
            return ok, reply, meta
        print("[loco4coco] exec failed for %s (%s)"
              % (kind, (meta or {}).get("error")))
    else:
        meta = {"transport": "skipped", "error": "no budget left for exec"}
        print("[loco4coco] skipping exec for %s: %ss of %ss budget used"
              % (kind, int(time.time() - started), budget))

    # Layer 3: COMPLETE, if it was not already the lead.
    if wants != "complete":
        ok3, reply3, meta3 = run_complete(cfg, prompt, kind, job_id=job_id,
                                          model=loc.get("complete_model"))
        if ok3:
            meta3["transport"] = "complete-fallback"
            return True, reply3, meta3

    # Layer 4 is the caller's business: it holds the archetype and its defaults,
    # so it can answer with no model at all. Returning not-ok is how it is asked.
    return False, "", (meta or {"transport": "none",
                                "error": "all layers failed"})


def base_context(cfg, state):
    vis = state.get("visitor") or {}
    ind = industry_name(cfg, vis.get("industry"))
    ctx = [
        f"You are CoCo, the Snowflake penguin, at the Snowflake World Tour "
        f"{cfg['event']['city']} booth.",
        "This is a five-minute activation. Be warm, brief and concrete.",
        "Never use bullet points, headings, markdown or emoji - your words are "
        "shown in a small speech bubble.",
        "You are already mid-conversation: do not open with a greeting and do "
        "not start by saying their name.",
    ]
    if vis.get("first_name"):
        ctx.append(f"You are speaking to {vis['first_name']}"
                   + (f" from {vis['company']}" if vis.get("company") else "")
                   + (f", in {ind}." if ind else "."))
    return ctx


def fill(tmpl, cfg, state, **extra):
    vis = state.get("visitor") or {}
    vals = {
        "first_name": vis.get("first_name") or "there",
        "company": vis.get("company") or "your organisation",
        "email": vis.get("email") or "you",
        "industry": industry_name(cfg, vis.get("industry")) or "your sector",
        "held": ", ".join(state.get("held") or []) or "nothing yet",
        "joined": ", ".join(state.get("joined") or []) or "nothing yet",
        # Captured at the letter stage, so it is already in hand by the time the
        # visitor reaches the library. Every location prompt can lean on it,
        # which is why the later questions can be shorter than they were.
        "problem": (vis.get("problem") or "").strip() or "not stated",
        "platforms": ", ".join(state.get("platforms") or []) or "not stated",
        # Home-stage sovereignty answers, so the reveal can reassure in-region.
        "company_country": (vis.get("company_country") or "").strip() or "not stated",
        "residency": {"country_only": "must stay in their country",
                      "eu": "must stay in the EU", "us_ok": "the US is acceptable too",
                      "unsure": "not yet decided"}.get(
                          vis.get("residency") or "", "not stated"),
    }
    vals.update(extra)
    out = tmpl
    for k, v in vals.items():
        out = out.replace("{" + k + "}", str(v))
    return out


# ------------------------------------------------------------------- industries

def industry_name(cfg, key):
    if not key:
        return ""
    return ((cfg.get("industries") or {}).get(key) or {}).get("name", "")


def infer_industry(cfg, company):
    """Keyword match on the company name. Cheap, instant, and always
    confirmed by the visitor with one click."""
    c = (company or "").lower()
    if not c:
        return "other"
    for key, hints in INDUSTRY_HINTS.items():
        for h in hints:
            if h in c:
                return key
    return "other"


def options_for(cfg, loc_id, state):
    """The checklist a location shows.

    The Marketplace reads the verified listing index rather than config, so
    the visitor is offered real listings with real providers - Tier 0
    agentic first if it finished in time, then Tier 1 live, then Tier 2
    curated. See listings_for().
    """
    loc = (cfg.get("locations") or {}).get(loc_id) or {}
    src = loc.get("source")
    if not src:
        return []
    ind = (state.get("visitor") or {}).get("industry") or "other"
    if src == "marketplace_index":
        return [{"id": r["global_name"], "label": r["title"],
                 "note": f"{r['provider']} \u00b7 {r['access']}",
                 "url": r["url"], "provider": r["provider"],
                 "access": r["access"]}
                for r in listings_for(cfg, ind, state.get("session_id"), state)]
    industries = cfg.get("industries") or {}
    block = industries.get(ind) or industries.get("other") or {}
    return block.get(src) or []


# ----------------------------------------------------------------- guides index

_guides_cache = None


def load_guides():
    """Parse guides-index.md into {archetype: (title, slug)} primary forks."""
    global _guides_cache
    if _guides_cache is not None:
        return _guides_cache
    try:
        rows = bctx.load(conn=coco_connection())[0].get("guides") or []
        prim = {}
        for r in rows:
            if r.get("is_primary") and r.get("archetype") not in prim:
                prim[r["archetype"]] = (r.get("title"), r.get("slug"))
        if prim:
            _guides_cache = prim
            return _guides_cache
    except Exception:
        pass
    out, current = {}, None
    try:
        with open(GUIDES_PATH, encoding="utf-8") as f:
            for line in f:
                m = re.match(r"^##\s+\d+\.\s+(\S+)", line)
                if m:
                    current = m.group(1).strip()
                    continue
                if current and line.startswith("|") and "`" in line:
                    cells = [c.strip() for c in line.strip().strip("|").split("|")]
                    if len(cells) < 2:
                        continue
                    slug = cells[1].strip("` ")
                    note = cells[2] if len(cells) > 2 else ""
                    if "Primary fork" in note and current not in out:
                        out[current] = (cells[0], slug)
    except OSError:
        pass
    _guides_cache = out
    return out


def guide_for(cfg, archetype):
    g = load_guides().get((archetype or "").strip())
    if not g:
        return "", ""
    base = (cfg.get("delivery") or {}).get(
        "guides_base", "https://www.snowflake.com/en/developers/guides/")
    return g[0], base + g[1] + "/"


# ------------------------------------------------- feature documentation links

_features_cache = None


def coco_connection():
    """The Snowflake connection name the booth is configured to use.

    Read from config rather than passed down: the context loaders are called
    from a dozen places, several of them before any request exists, and
    threading a connection through all of them buys nothing.
    """
    try:
        cfg = load_config()
        return ((cfg.get("coco") or {}).get("connection")
                or (cfg.get("snowflake") or {}).get("connection_name") or None)
    except Exception:
        return None


def load_features():
    """Parse feature-docs.md into {feature name: docs url}.

    This is a CLOSED list: the Workshop may only name features from here, which
    is what guarantees every feature in a blueprint carries a working link. A
    name outside the list is dropped rather than rendered bare - omitting one is
    better than shipping a guess in something a visitor keeps.
    """
    global _features_cache
    if _features_cache is not None:
        return _features_cache
    # Prefer the shared context (Snowflake, then the committed bundle) so a
    # laptop that never pulled this branch still gets the current closed list.
    # context.load() falls back to this same markdown, so the try below is the
    # last resort of a last resort rather than a duplicate path.
    try:
        rows = bctx.load(conn=coco_connection())[0].get("features") or []
        if rows:
            _features_cache = {r["name"]: r["docs_url"] for r in rows}
            return _features_cache
    except Exception:
        pass
    out = {}
    try:
        with open(FEATURES_PATH, encoding="utf-8") as f:
            for line in f:
                if not line.startswith("|") or "https://" not in line:
                    continue
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if len(cells) >= 2 and cells[1].startswith("https://"):
                    out[cells[0]] = cells[1]
    except OSError:
        pass
    _features_cache = out
    return out


def link_features(names):
    """Return [(name, url)] for names in the closed list, preserving order."""
    docs = load_features()
    seen, out = set(), []
    for n in (names or []):
        key = str(n).strip()
        if key in docs and key not in seen:
            seen.add(key)
            out.append((key, docs[key]))
    return out


ARCHETYPES_PATH = os.path.join(HERE, "archetypes.md")
_arch_cache = None


def load_archetypes():
    """Parse archetypes.md into {key: {features, first_step, pool}}.

    This is the precompute that makes the forge fast. Everything here is
    available with zero inference, so the blocking model call only has to choose
    an archetype and speak, and the background call returns consideration
    INDICES instead of four sentences of prose.
    """
    global _arch_cache
    if _arch_cache is not None:
        return _arch_cache
    out, cur = {}, None
    try:
        with open(ARCHETYPES_PATH, encoding="utf-8") as f:
            for line in f:
                line = line.rstrip("\n")
                if line.startswith("## "):
                    cur = line[3:].strip()
                    out[cur] = {"features": [], "first_step": "", "pool": []}
                    continue
                if not cur:
                    continue
                m = re.match(r"^\|\s*features\s*\|\s*(.+?)\s*\|$", line)
                if m:
                    out[cur]["features"] = [x.strip() for x in
                                            m.group(1).split(",") if x.strip()]
                    continue
                m = re.match(r"^\|\s*first_step\s*\|\s*(.+?)\s*\|$", line)
                if m:
                    out[cur]["first_step"] = m.group(1).strip()
                    continue
                m = re.match(r"^\d+\.\s+(.+)$", line)
                if m:
                    out[cur]["pool"].append(m.group(1).strip())
    except OSError:
        pass
    _arch_cache = out
    return out


def archetype_defaults(cfg, arche):
    """Everything we can say about an archetype without asking a model."""
    cat = load_archetypes()
    a = cat.get(arche) or cat.get("talk-to-my-data") or {}
    return {
        "features": [n for n, _u in link_features(a.get("features") or [])],
        "first_step": a.get("first_step") or "",
        "pool": a.get("pool") or [],
    }


# ------------------------------------------------------- marketplace listings

_market_cache = None


def load_marketplace():
    """Parse marketplace-index.md into {industry: [listing dicts]}.

    Only listings in this file may ever be named. SQL exposes a provider name
    for a small minority of listings and the public page is client-rendered, so
    provider names are recorded constants - see the file header.
    """
    global _market_cache
    if _market_cache is not None:
        return _market_cache
    try:
        rows = bctx.load(conn=coco_connection())[0].get("listings") or []
        if rows:
            grouped = {}
            for r in sorted(rows, key=lambda x: (x.get("industry") or "",
                                                 x.get("ordinal") or 0)):
                grouped.setdefault(r.get("industry"), []).append({
                    "title": r.get("title"), "provider": r.get("provider"),
                    "access": r.get("access"), "url": r.get("url"),
                    "global_name": r.get("global_name"),
                    "regions": r.get("regions") or "",
                })
            _market_cache = grouped
            return _market_cache
    except Exception:
        pass
    out, current = {}, None
    row_re = re.compile(r"^\|\s*\[(?P<title>.+?)\]\((?P<url>[^)]+)\)\s*\|"
                        r"\s*(?P<prov>[^|]+?)\s*\|\s*(?P<acc>[^|]+?)\s*\|"
                        r"\s*`(?P<gname>[^`]+)`\s*\|\s*(?P<regs>[^|]*)\|")
    try:
        with open(MARKET_PATH, encoding="utf-8") as f:
            for line in f:
                h = re.match(r"^##\s+([a-z_]+)\s*$", line.strip())
                if h:
                    current = h.group(1)
                    out.setdefault(current, [])
                    continue
                if not current:
                    continue
                m = row_re.match(line.strip())
                if m:
                    out[current].append({
                        "title": m.group("title").strip(),
                        "provider": m.group("prov").strip(),
                        "access": m.group("acc").strip(),
                        "url": m.group("url").strip(),
                        "global_name": m.group("gname").strip(),
                        "regions": m.group("regs").strip(),
                    })
    except OSError:
        pass
    _market_cache = out
    return out


_live_cache = {"at": 0, "rows": [], "error": ""}
_live_lock = threading.Lock()

# Tier 0: agentic marketplace search, fired once per visitor from _intake()
# (THE LETTER), racing THE LIBRARY. Single-visitor state, so a session_id
# guard is enough - no per-session dict needed. See listings_for().
_agentic_cache = {"session_id": None, "status": "idle", "rows": [], "at": 0}
_agentic_lock = threading.Lock()


def run_agentic_search(cfg, industry, problem):
    """Ask the marketplace-search skill for listings matched to what this
    visitor actually typed, not just their industry bucket.

    Deliberately does not go through run_exec(): that pushes live reasoning
    into state.reasoning for the visible turn on screen, and this call runs
    silently in the background while the visitor is elsewhere in the flow.
    Same NDJSON parsing, no UI side effects.

    Returns (rows, usage, ok, seconds). Never raises - every failure mode
    (CLI missing, timeout, bad JSON, empty result) comes back as ([], {}, False, t).
    """
    ac = (cfg.get("marketplace") or {}).get("agentic") or {}
    c = cfg.get("coco") or {}
    ind_name = industry_name(cfg, industry) or industry
    prompt = (
        f"Use the marketplace-search skill to find Snowflake Marketplace "
        f"listings relevant to this problem from a visitor in {ind_name}: "
        f"\"{problem}\". Return ONLY a JSON array (no prose, no code fence) "
        f"of objects with keys: title, provider, description, and "
        f"global_name if you know it.")
    cmd = [c.get("binary", "cortex"), "exec", prompt, "--format", "json",
           "--bypass", "--no-history"]
    conn = ac.get("connection") or (cfg.get("snowflake") or {}).get("connection_name")
    if conn:
        cmd += ["--connection", conn]
    timeout = float(ac.get("timeout_seconds", 70))
    started = time.time()
    final, usage, ok = "", {}, False
    try:
        p = subprocess.run(cmd, capture_output=True, text=True,
                           timeout=timeout, cwd=HERE)
        for raw in (p.stdout or "").splitlines():
            raw = raw.strip()
            if not raw:
                continue
            try:
                ev = json.loads(raw)
            except json.JSONDecodeError:
                continue
            if ev.get("type") == "result":
                final = (ev.get("result") or "").strip()
                usage = ev.get("usage") or {}
                ok = not ev.get("is_error")
    except subprocess.TimeoutExpired:
        pass
    except Exception:                                            # noqa: BLE001
        pass
    seconds = time.time() - started
    log_cost(cfg, "marketplace_agentic", seconds, usage, ok, "exec")
    rows = []
    if ok and final:
        m = re.search(r"\[.*\]", final, re.S)
        if m:
            try:
                data = json.loads(m.group(0))
                rows = [{"title": str(r.get("title") or "").strip(),
                         "provider": str(r.get("provider") or "Snowflake Marketplace").strip(),
                         "access": "Suggested by agentic search",
                         "url": LISTING_URL + str(r.get("global_name") or ""),
                         "global_name": str(r.get("global_name") or ""),
                         "regions": ""}
                        for r in data if isinstance(r, dict) and r.get("title")]
            except (json.JSONDecodeError, TypeError):
                rows = []
    return rows, usage, ok, seconds


def start_agentic_search(cfg, session_id, industry, problem):
    """Fire-and-forget: called from _intake() the moment THE LETTER is
    submitted. Runs on a daemon thread so the Letter response is never
    delayed. See run_agentic_search() and listings_agentic()."""
    ac = (cfg.get("marketplace") or {}).get("agentic") or {}
    if not ac.get("enabled") or not (problem or "").strip():
        return

    def job():
        with _agentic_lock:
            _agentic_cache.update(session_id=session_id, status="pending",
                                  rows=[], at=time.time())
        rows, _usage, ok, _s = run_agentic_search(cfg, industry, problem)
        with _agentic_lock:
            if _agentic_cache.get("session_id") != session_id:
                return              # a new visitor started while we ran
            _agentic_cache.update(
                status="ready" if (ok and rows) else "failed",
                rows=rows, at=time.time())

    threading.Thread(target=job, daemon=True).start()


def listings_agentic(cfg, session_id):
    """Tier 0 read path. Ready only if this visitor's call finished, matched
    the deny list Tier 1 already uses, and actually returned rows."""
    if not session_id:
        return []
    with _agentic_lock:
        if (_agentic_cache.get("session_id") != session_id
                or _agentic_cache.get("status") != "ready"):
            return []
        rows = list(_agentic_cache.get("rows") or [])
    mk = cfg.get("marketplace") or {}
    deny = [w.lower() for w in (mk.get("exclude") or [])]
    return [r for r in rows
            if not any(d in (r.get("title") or "").lower() for d in deny)]


def region_short(cfg):
    return ((cfg.get("event") or {}).get("region") or "").strip().split(".")[-1]


def refresh_live_listings(cfg, force=False):
    """Pull the real Marketplace catalogue once, then serve every visitor from memory.

    No private preview needed: SHOW AVAILABLE LISTINGS piped through RESULT_SCAN
    is enough. Measured on PG_LONDON: 4327 listings, 799 available in
    AWS_EU_WEST_2, 188 of those importable. One query per server start rather
    than one per visitor, so the stall stays instant.

    Note SHOW ... LIKE does NOT filter on title, which is why the title match is
    done here in SQL and not in the SHOW.
    """
    mk = cfg.get("marketplace") or {}
    ttl = float(mk.get("cache_minutes", 120)) * 60
    with _live_lock:
        if not force and _live_cache["rows"] and time.time() - _live_cache["at"] < ttl:
            return _live_cache["rows"], ""
        region = region_short(cfg)
        if not region:
            _live_cache["error"] = "event.region is not set"
            return [], _live_cache["error"]
        try:
            cur = sf_conn(cfg).cursor()
            try:
                cur.execute("SHOW AVAILABLE LISTINGS")
                cur.execute(
                    'SELECT "title", "global_name", "profile", "regions" '
                    "FROM TABLE(RESULT_SCAN(LAST_QUERY_ID(-1))) "
                    'WHERE "regions" LIKE %s AND "is_ready_for_import" = \'true\'',
                    (f"%{region}%",))
                rows = [{"title": t, "global_name": g, "profile": p,
                         "regions": r,
                         # is_ready_for_import was already filtered to true
                         "access": "Ready to import",
                         "url": LISTING_URL + g}
                        for (t, g, p, r) in cur.fetchall()]
            finally:
                cur.close()
        except Exception as e:                                   # noqa: BLE001
            _live_cache["error"] = str(e)[:200].replace("\n", " ")
            return [], _live_cache["error"]
        _live_cache.update({"at": time.time(), "rows": rows, "error": ""})
        return rows, ""


def listings_live(cfg, industry):
    """Match the live catalogue against this industry's keyword avenues.

    Pinned listings (verified by global_name) come first so a booth always leads
    with something known-good, then keyword matches fill the rest.
    """
    rows, err = refresh_live_listings(cfg)
    if not rows:
        return [], err or "no live listings"
    mk = cfg.get("marketplace") or {}
    words = [w.lower() for w in
             ((mk.get("industry_keywords") or {}).get(industry)
              or (mk.get("industry_keywords") or {}).get("other") or [])]
    pins = (mk.get("pinned") or {}).get(industry) or []
    deny = [w.lower() for w in (mk.get("exclude") or [])]

    def blocked(title):
        t = (title or "").lower()
        return any(d in t for d in deny)

    by_gn = {r["global_name"]: r for r in rows}
    names = mk.get("provider_names") or {}

    def named(r):
        # SHOW returns an opaque profile id, not a display name, so a known
        # provider is named from config and anything else is labelled honestly.
        r = dict(r)
        r["provider"] = names.get(r.get("profile") or "", "Snowflake Marketplace")
        return r

    # Geographic relevance. MEASURED: a public-sector visitor at a UK stand was
    # offered "India Economic Monitor" and "Weather Data - All International
    # Postal Codes" because keyword matches were taken in raw catalogue order
    # with no notion of where the event is. Config-driven so moving the event to
    # another country is a config change, not a code change.
    geo = (mk.get("geo") or {})
    prefer = [w.lower() for w in (geo.get("prefer") or [])]
    demote = [w.lower() for w in (geo.get("demote") or [])]
    # Per-industry veto: "postcode" is a legitimate public-sector keyword but it
    # matched "Postcode Sector Weather Forecasts", which is how three of the six
    # public-sector picks ended up being Met Office weather products.
    avoid = [w.lower() for w in
             ((mk.get("industry_avoid") or {}).get(industry) or [])]

    def relevance(r):
        """Higher is better. (geo_score, keyword_hits) decides the order."""
        hay = ((r.get("title") or "") + " " +
               (names.get(r.get("profile") or "", ""))).lower()
        g = sum(2 for w in prefer if w in hay) - sum(3 for w in demote if w in hay)
        k = sum(1 for w in words if w in hay)
        return (g, k)

    out, seen = [], set()
    for gn in pins:
        r = by_gn.get(gn)
        if r:
            out.append(named(r))
            seen.add(gn)
    cands = []
    for r in rows:
        if r["global_name"] in seen:
            continue
        t = (r["title"] or "").lower()
        if not any(w in t for w in words) or blocked(t):
            continue
        if any(a in t for a in avoid):
            continue
        cands.append(r)
    # Deterministic: sort by relevance, then title, so the same catalogue always
    # produces the same stall. Ties used to fall out in whatever order SHOW
    # happened to return, which made the booth unpredictable between visitors.
    cands.sort(key=lambda r: (relevance(r), (r.get("title") or "").lower()),
               reverse=True)
    for r in cands:
        out.append(named(r))
        seen.add(r["global_name"])
    return out, ""


def listings_for(cfg, industry, session_id=None, state=None):
    """Tier chain: agentic first if it's ready, live catalogue next, curated
    index as the booth-safe net.

    Tier 0 agentic  - cortex exec calling the marketplace-search skill,
                     started from _intake() at THE LETTER and racing THE
                     LIBRARY. Personalised to what this visitor actually
                     typed, not just their industry bucket. Used only if it
                     finished in time - see listings_agentic().
    Tier 1 live     - real, current, region-filtered, importable.
    Tier 2 curated  - marketplace-index.md, used when live is cold, fails, or
                     is too thin to fill a stall. Offline safe and quality
                     checked.

    "Agentic search on the Snowflake Marketplace" (PrPr) was verified working
    on PG_LONDON on 2026-08-23 via both Snowsight's Discover tab AND
    `cortex exec` calling the same marketplace-search skill from the CLI -
    the latter is what Tier 0 actually calls. It never blocks a visitor: if
    it hasn't finished, errored, or `marketplace.agentic.enabled` is false,
    this falls straight through to Tier 1 then Tier 2 exactly as before.
    """
    loc = ((cfg.get("locations") or {}).get("marketplace") or {})
    mk = cfg.get("marketplace") or {}
    want = int(mk.get("min_live_results", 3))
    agentic = listings_agentic(cfg, session_id)
    if len(agentic) >= want:
        _live_cache["tier"] = "agentic"
        return agentic[:max(want, 6)]
    if (loc.get("discovery") or "manual") == "live":
        live, err = listings_live(cfg, industry)
        if len(live) >= want:
            _live_cache["tier"] = "live"
            return live[:max(want, 6)]
        print(f"[loco4coco] live marketplace thin for {industry} "
              f"({len(live)} hits{', ' + err if err else ''}), using curated index")
    _live_cache["tier"] = "curated"
    return listings_curated(cfg, industry, state)


def _problem_tokens(state):
    """Content words from what the visitor actually typed, for matching titles."""
    vis = ((state or {}).get("visitor") or {})
    txt = " ".join([str(vis.get("problem") or ""),
                     str(vis.get("company") or "")]).lower()
    stop = {"the", "and", "our", "for", "with", "from", "into", "that", "this",
            "have", "has", "are", "was", "we", "us", "to", "of", "in", "on",
            "it", "is", "a", "an", "by", "at", "be", "do", "not", "but",
            "data", "snowflake", "want", "need", "lot", "lots", "time"}
    out = set()
    for w in re.split(r"[^a-z0-9]+", txt):
        if len(w) > 3 and w not in stop:
            out.add(w)
            if w.endswith("s"):
                out.add(w[:-1])
    return out


def listings_curated(cfg, industry, state=None):
    """Curated listings, ordered by what this visitor typed, filtered to region.

    Two visitors in the same industry used to get an IDENTICAL stall: this
    returned a fixed six in file order, and since `marketplace.agentic.enabled`
    is false and `locations.marketplace.discovery` is "manual", this is the only
    tier that ever runs. So the stall was a function of the industry bucket
    alone, and the same six appeared over and over.

    Two changes, neither of which invents a listing - every row here is a real,
    checked Marketplace entry and that discipline is not worth trading for
    variety:

      1. The pool is widened by BORROWING from other buckets. A visitor whose
         problem mentions weather, or property, or company registrations is
         better served by the real listing that matches it than by the sixth-best
         pick from their own sector. Their own bucket still leads, via a base
         score.
      2. Ties rotate on the session id, so consecutive visitors in one bucket do
         not see the same order.

    The region filter is not cosmetic: handing a London visitor a us-east-1-only
    share sends them somewhere they cannot go.
    """
    market = load_marketplace()
    own = market.get(industry) or market.get("other") or []
    toks = _problem_tokens(state)
    seen, scored = set(), []
    for bucket, rows in [(industry, own)] + sorted(market.items()):
        for r in rows:
            gn = r.get("global_name") or r.get("title")
            if gn in seen:
                continue
            seen.add(gn)
            hay = (str(r.get("title") or "") + " "
                   + str(r.get("provider") or "")).lower()
            hits = sum(1 for t in toks if t in hay)
            base = 2 if r in own else 0
            scored.append((base + hits * 3, hits, r))
    # Stable per-session shuffle for the ties, so the stall changes visitor to
    # visitor without becoming random - the same session always sees the same
    # stall, which matters when the panel is reopened.
    salt = str(((state or {}).get("session_id")) or "")
    def key(item):
        s, hits, r = item
        h = hashlib.md5((salt + str(r.get("global_name") or "")).encode()).hexdigest()
        return (-s, h)
    rows = [r for _, _, r in sorted(scored, key=key)]
    region = ((cfg.get("event") or {}).get("region") or "").strip()
    if not region:
        return rows[:6]
    short = region.split(".")[-1]
    keep = []
    for r in rows:
        regs = (r.get("regions") or "").strip()
        # "ALL" means every region. MEASURED: this was being substring-matched
        # like a region list, so it matched nothing and silently dropped the
        # listing - which is how public sector shipped 5 picks instead of 6 and
        # lost Ordnance Survey Boundary Line, one of the best UK public-sector
        # datasets on the Marketplace.
        if regs.upper() == "ALL":
            keep.append(r)
            continue
        # A truncated "+N" list means we cannot prove absence, so keep it.
        if short in regs or "+" in regs:
            keep.append(r)
    return keep[:6]


# --------------------------------------------------------------- the CoCo prompt

def build_coco_prompt(cfg, state):
    """The thing they paste into CoCo on their own free trial.

    Six parts, per prompt-builder.md. Built deterministically from what they
    actually told us: it must never invent table or column names, so it tells
    CoCo to inspect and ask instead.
    """
    vis = state.get("visitor") or {}
    poc = state.get("poc") or {}
    ind = industry_name(cfg, vis.get("industry")) or "our sector"
    held = state.get("held") or []
    joined = state.get("joined") or []
    listings = state.get("joined_listings") or []

    lines = [
        f"You are helping me build a proof of concept in Snowflake. I work in "
        f"{ind}" + (f" at {vis.get('company')}" if vis.get("company") else "") + ".",
        "",
        "THE PROBLEM",
        (vis.get("problem") or "").strip()
        or "(I described this at the booth but did not write it down.)",
        "",
        "THE DATA",
        "I have not loaded anything yet. The data I hold is:",
    ]
    lines += [f"  - {h}" for h in (held or ["(to be confirmed)"])]
    plats = state.get("platforms") or []
    if plats:
        lines.append("It is sitting on: " + ", ".join(plats)
                     + ". Tell me the shortest route to get it into Snowflake "
                       "before you build anything.")
    if listings:
        lines.append("I also want to attach these Snowflake Marketplace listings:")
        for r in listings:
            lines.append(f"  - {r['title']} (provider: {r['provider']}, "
                         f"listing {r['global_name']})")
        extra = [j for j in joined if j not in {r["title"] for r in listings}]
        for j in extra:
            lines.append(f"  - {j} (I will need to find a suitable listing)")
    elif joined:
        lines.append("I also want to join these categories of Marketplace data:")
        lines += [f"  - {j}" for j in joined]
    lines += [
        "",
        "BUILD THIS",
        poc.get("poc_name") or "A working proof of concept",
    ]
    if poc.get("summary"):
        lines.append(poc["summary"])
    feats = [n for n, _u in link_features(poc.get("features"))]
    if feats:
        lines.append("Use these Snowflake features: " + ", ".join(feats) + ".")

    title, url = poc.get("guide_title"), poc.get("guide_url")
    if title and url:
        lines += ["", "START FROM", f"Follow this guide as the base: {title}", url]

    lines += [
        "",
        "CONSTRAINTS",
        "Do not invent table or column names. Inspect what exists first, and "
        "ask me when something is ambiguous.",
        "Work in small steps and show me the SQL before you run anything that "
        "creates or changes objects.",
        "Keep everything in one database and schema so it is easy to drop afterwards.",
        "",
        "FIRST STEP",
        poc.get("first_step") or "Tell me what you need from me to get started, "
        "then set up the database and schema.",
    ]
    return "\n".join(lines)


# ------------------------------------------------------------------- blueprint

# The blueprint email body (blueprint_html) was removed with the outbox
# on 2026-08-27: nothing composes an email any more, the QR to the
# presigned .docx is the delivery, and the on-screen read surface is
# blueprint_page below.


INTEGRATION_PATHS = {
    "Microsoft / Azure": (
        "Openflow has a first-party connector for Azure Blob Storage and "
        "SQL Server. For Fabric or OneLake, register the Iceberg tables through "
        "a catalog integration and query them in place - no copy.",
        "https://docs.snowflake.com/en/user-guide/data-load-azure"),
    "AWS": (
        "Point an external stage at the S3 bucket with a storage integration, "
        "then Snowpipe for continuous load. If the data is already Iceberg in "
        "Glue, use a catalog integration and leave it where it is.",
        "https://docs.snowflake.com/en/user-guide/data-load-s3"),
    "Google Cloud": (
        "A storage integration over the GCS bucket plus an external stage. "
        "BigQuery data moves cleanly as Parquet exported to GCS, or through "
        "Openflow if you need it on a schedule.",
        "https://docs.snowflake.com/en/user-guide/data-load-gcs"),
    "Oracle": (
        "Openflow's Oracle connector does change data capture, so you get an "
        "ongoing replica rather than a nightly dump. Start with the handful of "
        "tables the proof of concept actually reads.",
        "https://other-docs.snowflake.com/en/connectors"),
    "SAP": (
        "Either the SAP connector for Snowflake, or SAP Business Data Cloud "
        "sharing the data as Iceberg that Snowflake reads without a copy. The "
        "second route is usually faster to stand up.",
        "https://docs.snowflake.com/en/user-guide/tables-iceberg"),
    "On-premise / our own servers": (
        "Openflow can run inside your network and push out, so nothing has to "
        "be exposed inbound. For a first proof of concept, a one-off bulk load "
        "of a representative extract is usually enough.",
        "https://docs.snowflake.com/en/user-guide/data-load-local-file-system"),
    "SaaS apps (Salesforce, Workday, etc.)": (
        "Openflow has connectors for the common SaaS sources, and the "
        "Marketplace carries some of them as ready-made shares. Check the "
        "Marketplace first - it is the cheaper answer when it exists.",
        "https://other-docs.snowflake.com/en/connectors"),
    "Already in Snowflake": (
        "Nothing to move. Point the proof of concept at the existing tables "
        "and spend the saved time on the model and the interface instead.",
        "https://docs.snowflake.com/en/guides-overview-queries"),
    "Not sure yet": (
        "Worth ten minutes with whoever owns the source before you build. The "
        "answer changes the effort more than any other decision here.",
        "https://docs.snowflake.com/en/guides-overview-loading-data"),
}


def normalise_platforms(cfg, plats):
    """Resolve a raw platform selection into a predictable, non-contradictory set.

    MEASURED before this existed: of the 36 possible pairs, 16 were logically
    contradictory and 10 printed "Openflow" twice. Selecting "Already in
    Snowflake" alongside Oracle produced a blueprint that said "Nothing to move"
    AND gave move instructions; "Not sure yet" plus AWS said "go and ask whoever
    owns the source" next to a concrete S3 route. At a Snowflake-branded event
    the document has to be defensible, so the contradiction is resolved here
    rather than left to chance.

    Rules, in order:
      1. Keep only chips the config actually offers, de-duplicated.
      2. Restore config order, so the same taps always produce the same document
         regardless of the order they were tapped in.
      3. An `exclusive` chip ("nothing to move" / "go and ask") survives only if
         NOTHING else was picked. A named source is actionable and wins.
      4. Cap at `max_routes` so the ingestion section stays a plan, not a list.

    The UI enforces 3 on tap as well; this is the backstop, because the document
    is the thing that leaves the stand.
    """
    cfgp = (cfg.get("platforms") or {})
    offered = list(cfgp.get("options") or [])
    excl = set(cfgp.get("exclusive") or [])
    cap = int(cfgp.get("max_routes") or 4)

    seen, kept = set(), []
    for p in (plats or []):
        if p in offered and p not in seen:
            seen.add(p)
            kept.append(p)
    kept.sort(key=offered.index)                  # deterministic output

    named = [p for p in kept if p not in excl]
    if named:
        kept = named                              # a real source beats "not sure"
    else:
        kept = kept[:1]                           # at most one exclusive answer
    return kept[:cap]


def sovereignty_lines(cfg, state):
    """The four in-region pillars, personalised to the visitor's residency.

    Precompute, not generation: the prose is ours (config.sovereignty.pillars),
    so it is consistent between visitors and cannot name a control that does not
    exist. Returns [] when we have nothing to personalise against, so a visitor
    who skipped the question does not get a generic lecture.
    """
    vis = state.get("visitor") or {}
    residency = (vis.get("residency") or "").strip().lower()
    country = (vis.get("company_country") or "").strip()
    if not residency and not country and not state.get("platforms"):
        return []
    region = {"country_only": country or "your country",
              "eu": "the EU",
              "us_ok": "your chosen region",
              "unsure": "your chosen region"}.get(residency, "your chosen region")
    pillars = (cfg.get("sovereignty") or {}).get("pillars") or {}
    order = ["data", "models", "marketplace", "governance"]
    return [pillars[k].replace("{region}", region) for k in order if pillars.get(k)]


def integration_paths(state):
    """The route into Snowflake for each platform the visitor named."""
    out = []
    for p in normalise_platforms(load_config(), state.get("platforms")):
        hit = INTEGRATION_PATHS.get(p)
        if hit:
            out.append((p, hit[0], hit[1]))
        else:
            out.append((p, "Check the connector catalogue for this source, or "
                        "land an extract on a stage as Parquet to begin with.",
                        "https://other-docs.snowflake.com/en/connectors"))
    return out


# The .docx is a good thing to forward to a colleague and a poor thing to read on
# a phone: Word on iOS wants a download and an app switch before anyone sees a
# word of it, and a QR code at a booth is scanned by a phone every single time.
# So the same state renders twice - a page for reading now, a document for
# keeping. Mobile-first is not a nicety here, it is the primary case.
BLUEPRINT_CSS = """
:root { color-scheme:light; }
* { box-sizing:border-box; }
body { margin:0; padding:0 18px 64px;
  font:16px/1.6 -apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Helvetica,
       Arial,sans-serif;
  color:#1A242E; background:#F4F8FB;
  -webkit-text-size-adjust:100%; }
.wrap { max-width:34rem; margin:0 auto; }
header { padding:26px 0 18px; }
.mark { display:block; margin-bottom:14px; }
h1 { font-size:1.55rem; line-height:1.25; margin:0 0 6px; color:#11567F;
  letter-spacing:-.01em; }
.for { color:#55707F; font-size:.95rem; margin:0; }
.lede { font-size:1.06rem; color:#243642; margin:16px 0 0; }
section { background:#fff; border:1px solid #DCE7EF; border-radius:12px;
  padding:16px 16px 18px; margin:14px 0; }
h2 { font-size:.78rem; letter-spacing:.10em; text-transform:uppercase;
  color:#29B5E8; margin:0 0 10px; }
ul { margin:0; padding-left:1.15rem; }
li { margin:0 0 9px; }
li b { color:#11567F; }
.note { display:block; color:#5B7382; font-size:.88rem; margin-top:2px; }
a { color:#0B7FB3; overflow-wrap:anywhere; }
/* Tap targets, not links: a 14px underline on a phone is a coin toss. */
.btn { display:block; text-align:center; text-decoration:none;
  padding:15px 18px; border-radius:10px; font-weight:700; font-size:1rem;
  background:#29B5E8; color:#08222E; border:none; width:100%;
  cursor:pointer; font-family:inherit; }
.btn.alt { background:#fff; color:#11567F; border:2px solid #29B5E8;
  margin-top:10px; }
pre { margin:0; padding:14px; background:#0E1F2B; color:#DDEBF4;
  border-radius:10px; font:13px/1.5 ui-monospace,SFMono-Regular,Menlo,monospace;
  white-space:pre-wrap; overflow-wrap:anywhere; max-height:15rem;
  overflow-y:auto; }
ol { margin:0; padding-left:1.3rem; }
footer { text-align:center; color:#7B909D; font-size:.82rem; padding:22px 0 0; }
@media (prefers-color-scheme:dark) {
  body { background:#0B1620; color:#DCE8F0; }
  section { background:#12222E; border-color:#20384A; }
  h1,li b { color:#7FD3F2; } .for,.note { color:#8FA8B8; }
  .lede { color:#C9DCE8; } a { color:#7FD3F2; }
  .btn.alt { background:transparent; color:#7FD3F2; }
}
"""

SNOWFLAKE_SVG = (
    '<svg class="mark" width="34" height="34" viewBox="-50 -50 100 100" '
    'aria-hidden="true">'
    + "".join(
        '<g transform="rotate(%d)">'
        '<path d="M15 -4 L36 -4 L36 -15 L50 0 L36 15 L36 4 L15 4 Z" '
        'fill="#29B5E8"/></g>' % (i * 60) for i in range(6))
    + '<path d="M0 -15 L15 0 L0 15 L-15 0 Z" fill="#29B5E8"/>'
      '<path d="M0 -6 L6 0 L0 6 L-6 0 Z" fill="#F4F8FB"/></svg>')


def blueprint_page(cfg, state):
    """The blueprint as a standalone page, written to a temp file.

    Same data as the .docx, different medium. Three things it can do that the
    document cannot: render instantly on the phone that scanned the QR, make
    every doc link tappable, and put a copy button on the CoCo prompt - which is
    the one part of the blueprint the visitor has to move somewhere else.
    """
    e = html.escape
    vis = state.get("visitor") or {}
    poc = state.get("poc") or {}
    d = cfg.get("delivery") or {}
    held = state.get("held") or []
    joined = state.get("joined") or []
    listings = state.get("joined_listings") or []
    prompt = build_coco_prompt(cfg, state)
    signup = d.get("signup_url", "")
    city = ((cfg.get("event") or {}).get("city") or "")

    P = []
    a = P.append
    a("<!doctype html><html lang=\"en\"><head><meta charset=\"utf-8\">")
    a("<meta name=\"viewport\" content=\"width=device-width,"
      "initial-scale=1\">")
    a("<title>%s</title>" % e(poc.get("poc_name") or "Your Snowflake POC"))
    a("<style>%s</style></head><body><div class=\"wrap\">" % BLUEPRINT_CSS)

    a("<header>" + SNOWFLAKE_SVG)
    a("<h1>%s</h1>" % e(poc.get("poc_name") or "Your proof of concept"))
    who = e(vis.get("first_name") or "you")
    if vis.get("company"):
        who += " at " + e(vis["company"])
    a("<p class=\"for\">For %s%s</p>" % (
        who, " &middot; Snowflake World Tour " + e(city) if city else ""))
    if poc.get("summary"):
        a("<p class=\"lede\">%s</p>" % e(poc["summary"]))
    a("</header>")

    # The prompt goes FIRST on the page. It is the only thing that has to travel
    # somewhere else, and burying it under four reference sections on a phone is
    # how it gets lost.
    a("<section><h2>Start here</h2><ol>")
    a("<li>Start a free Snowflake trial%s.</li>"
      % (" at <a href=\"%s\">%s</a>" % (e(signup), e(signup)) if signup else ""))
    a("<li>Open Cortex Code.</li>")
    a("<li>Paste the prompt below as your first message.</li></ol>")
    a("<pre id=\"p\">%s</pre>" % e(prompt))
    a("<button class=\"btn\" id=\"c\" style=\"margin-top:12px\">"
      "Copy the prompt</button>")
    if signup:
        a("<a class=\"btn alt\" href=\"%s\">Start a free trial</a>" % e(signup))
    a("</section>")

    if poc.get("first_step"):
        a("<section><h2>Your first step</h2><p style=\"margin:0\">%s</p>"
          "</section>" % e(poc["first_step"]))

    if held:
        a("<section><h2>Data you already hold</h2><ul>")
        for h in held:
            a("<li>%s</li>" % e(h))
        a("</ul></section>")

    paths = integration_paths(state)
    if paths:
        a("<section><h2>Getting that data into Snowflake</h2><ul>")
        for plat, howto, url in paths:
            a("<li><b>%s</b><span class=\"note\">%s</span>"
              "<a href=\"%s\">%s</a></li>" % (e(plat), e(howto), e(url),
                                             e(url)))
        a("</ul></section>")

    sov = sovereignty_lines(cfg, state)
    if sov:
        a("<section><h2>Sovereignty and security</h2><ul>")
        for line in sov:
            a("<li>%s</li>" % e(line))
        a("</ul></section>")

    if listings or joined:
        a("<section><h2>Attach from the Marketplace</h2><ul>")
        named = set()
        for r in listings:
            named.add(r["title"])
            a("<li><a href=\"%s\"><b>%s</b></a>"
              "<span class=\"note\">%s &middot; %s</span></li>"
              % (e(r["url"]), e(r["title"]), e(r["provider"]), e(r["access"])))
        for j in joined:
            if j not in named:
                a("<li><b>%s</b><span class=\"note\">Search the Marketplace "
                  "for a provider</span></li>" % e(j))
        a("</ul></section>")

    feats = link_features(poc.get("features") or [])
    if feats:
        a("<section><h2>Features you will use</h2><ul>")
        for name, docurl in feats:
            a("<li><a href=\"%s\"><b>%s</b></a></li>" % (e(docurl), e(name)))
        a("</ul></section>")

    cons = poc.get("considerations") or []
    if cons:
        a("<section><h2>Worth thinking about</h2><ul>")
        for c in cons:
            a("<li>%s</li>" % e(str(c)))
        a("</ul></section>")

    problem = (vis.get("problem") or "").strip()
    if problem:
        a("<section><h2>The problem you described</h2>"
          "<p style=\"margin:0\">%s</p></section>" % e(problem))

    a("<footer>Built with you at the Snowflake booth.</footer>")
    a("</div><script>")
    # No inline handler and no remote script: this file gets served from a
    # presigned URL, so it has to be self-contained and boring.
    a("document.getElementById('c').addEventListener('click',function(){"
      "var t=document.getElementById('p').textContent;"
      "var b=this;navigator.clipboard&&navigator.clipboard.writeText(t)"
      ".then(function(){b.textContent='Copied';"
      "setTimeout(function(){b.textContent='Copy the prompt';},1800);});});")
    a("</script></body></html>")

    name = re.sub(r"[^A-Za-z0-9]+", "-",
                  (vis.get("first_name") or "visitor")).strip("-") or "visitor"
    path = os.path.join(tempfile.gettempdir(),
                        "loco4coco-%s-%d.html" % (name, int(time.time())))
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(P))
    return path


def blueprint_docx(cfg, state):
    """Generate a real .docx with python-docx. Returns a path, or "" if the
    library is missing - a missing document must not block the email."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        return ""
    vis = state.get("visitor") or {}
    poc = state.get("poc") or {}
    doc = Document()

    # Snowflake brand, not the python-docx default. The template ships Times New
    # Roman, which is the first thing a visitor notices about a document they
    # keep. Inter is the Snowflake typeface; the fallbacks cover a machine that
    # does not have it, and eastasia must be set too or Word substitutes a serif.
    BRAND_BLUE = RGBColor(0x29, 0xB5, 0xE8)
    HEADING_BLUE = RGBColor(0x11, 0x56, 0x7F)
    BODY_INK = RGBColor(0x1A, 0x24, 0x2E)

    def set_font(style_obj, name, size=None, color=None, bold=None):
        f = style_obj.font
        f.name = name
        if size is not None:
            f.size = size
        if color is not None:
            f.color.rgb = color
        if bold is not None:
            f.bold = bold
        rpr = style_obj.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), name)

    FONT = "Inter"
    set_font(doc.styles["Normal"], FONT, Pt(10.5), BODY_INK)
    for name, size in (("Title", Pt(24)), ("Heading 1", Pt(15)),
                       ("Heading 2", Pt(12.5)), ("List Bullet", Pt(10.5))):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        colour = BRAND_BLUE if name == "Title" else (
            HEADING_BLUE if name.startswith("Heading") else BODY_INK)
        set_font(st, FONT, size, colour, True if name != "List Bullet" else None)
        st.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        st.paragraph_format.space_after = Pt(4)

    h = doc.add_heading(poc.get("poc_name") or "Your proof of concept", level=0)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
        run.font.name = FONT
    doc.add_paragraph(
        f"Prepared for {vis.get('first_name') or 'you'}"
        + (f" at {vis.get('company')}" if vis.get("company") else "")
        + f" - Snowflake World Tour {cfg['event']['city']}.")
    if poc.get("summary"):
        doc.add_paragraph(poc["summary"])

    def bullets(title, items):
        if not items:
            return
        doc.add_heading(title, level=2)
        for i in items:
            doc.add_paragraph(str(i), style="List Bullet")

    bullets("Data you already hold", state.get("held") or [])

    paths = integration_paths(state)
    if paths:
        doc.add_heading("Getting that data into Snowflake", level=2)
        for plat, how, url in paths:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(plat).bold = True
            p.add_run(f"\n{how}\n{url}")

    bullets("Sovereignty and security", sovereignty_lines(cfg, state))

    problem = ((state.get("visitor") or {}).get("problem") or "").strip()
    if problem:
        doc.add_heading("The problem you described", level=2)
        doc.add_paragraph(problem)

    listings = state.get("joined_listings") or []
    joined = state.get("joined") or []
    if listings or joined:
        doc.add_heading("To attach from the Snowflake Marketplace", level=2)
        named = set()
        for r in listings:
            named.add(r["title"])
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(r["title"]).bold = True
            p.add_run(f"\n{r['provider']} \u00b7 {r['access']}\n{r['url']}")
        for j in joined:
            if j not in named:
                doc.add_paragraph(
                    f"{j} - search the Marketplace for a provider",
                    style="List Bullet")

    feats = link_features((poc.get("features") or []))
    if feats:
        doc.add_heading("Snowflake features you will use", level=2)
        for name, docurl in feats:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(name).bold = True
            p.add_run(f"\n{docurl}")

    signup = (cfg.get("delivery") or {}).get("signup_url", "")
    doc.add_heading("Before you start", level=2)
    doc.add_paragraph(
        "Start a free Snowflake trial" + (f" at {signup}" if signup else "")
        + ", then open Cortex Code.", style="List Number")
    doc.add_paragraph(
        "Copy the prompt in the next section and paste it in as your first "
        "message.", style="List Number")

    doc.add_heading("Paste this into Cortex Code to begin", level=2)
    p = doc.add_paragraph()
    run = p.add_run(build_coco_prompt(cfg, state))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), "Consolas")

    if poc.get("guide_title"):
        doc.add_heading("Start from this guide", level=2)
        doc.add_paragraph(poc["guide_title"])
        if poc.get("guide_url"):
            doc.add_paragraph(poc["guide_url"])
    cons = poc.get("considerations") or []
    if cons:
        doc.add_heading("Considerations", level=2)
        doc.add_paragraph("Worth thinking about before you start:")
        for c in cons:
            doc.add_paragraph(str(c), style="List Bullet")

    safe = re.sub(r"[^A-Za-z0-9]+", "-", (vis.get("first_name") or "visitor")).strip("-")
    path = os.path.join(tempfile.gettempdir(),
                        f"loco4coco-{safe}-{int(time.time())}.docx")
    doc.save(path)
    return path


# ----------------------------------------------------------------- snow helpers

def snow_sql(cfg, sql):
    """Run one statement via the snow CLI. Avoids adding a connector
    dependency to what is otherwise a stdlib server."""
    conn = (cfg.get("snowflake") or {}).get("connection_name")
    cmd = ["snow", "sql", "-q", sql, "--format", "json"]
    if conn:
        cmd += ["-c", conn]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
        if r.returncode != 0:
            return None, (r.stderr or r.stdout or "").strip()[:300]
        return json.loads(r.stdout or "[]"), ""
    except FileNotFoundError:
        return None, "snow CLI not found on PATH"
    except (json.JSONDecodeError, subprocess.TimeoutExpired) as e:
        return None, str(e)[:300]


def stage_and_presign(cfg, local_path):
    """Upload the .docx and return a presigned download URL."""
    d = cfg.get("delivery") or {}
    stage = d.get("stage") or ""
    if not stage or not local_path:
        return "", "no stage configured"
    conn = (cfg.get("snowflake") or {}).get("connection_name")
    cmd = ["snow", "stage", "copy", local_path, stage, "--overwrite"]
    if conn:
        cmd += ["-c", conn]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if r.returncode != 0:
            return "", (r.stderr or r.stdout or "").strip()[:300]
    except FileNotFoundError:
        return "", "snow CLI not found on PATH"
    except subprocess.TimeoutExpired:
        return "", "stage upload timed out"

    name = os.path.basename(local_path)
    secs = int(d.get("presign_seconds", 604800))
    rows, err = snow_sql(
        cfg, f"SELECT GET_PRESIGNED_URL({stage}, '{name}', {secs}) AS URL")
    if not rows:
        return "", err or "presign returned nothing"
    row = rows[0] if isinstance(rows, list) else rows
    if isinstance(row, list):
        row = row[0] if row else {}
    return (row.get("URL") or row.get("url") or ""), ""


def log_session(cfg, state):
    """Write the visitor's row. Returns (ok, error).

    Escaping is done ONCE, by handing Snowflake a single JSON document and
    letting it do the typing. Building 27 quoted literals by hand is how the
    previous version ended up writing to columns that did not exist and failing
    silently four times in a row.
    """
    sf = cfg.get("snowflake") or {}
    tbl = f"{sf.get('database')}.{sf.get('schema')}.{sf.get('sessions_table', 'SESSIONS')}"
    vis = state.get("visitor") or {}
    poc = state.get("poc") or {}
    turns = state.get("turns") or []
    qa = state.get("qa") or {}

    started = state.get("started_at") or time.time()
    payload = {
        "SESSION_ID": state.get("session_id") or "",
        "EVENT_CITY": (cfg.get("event") or {}).get("city", ""),
        "LANGUAGE_CODE": (cfg.get("event") or {}).get("language", ""),
        "FIRST_NAME": vis.get("first_name", ""),
        "COMPANY": vis.get("company", ""),
        "INDUSTRY": industry_name(cfg, vis.get("industry")),
        # The label is what a human reads in a report; the key is what
        # groups reliably. Storing only the label meant "Something else"
        # and a renamed industry were indistinguishable after the fact.
        "INDUSTRY_KEY": vis.get("industry") or "",
        "DATA_HELD": state.get("held") or [],
        "MARKETPLACE_JOINED": state.get("joined") or [],
        "POC_ARCHETYPE": poc.get("archetype", ""),
        "POC_NAME": poc.get("poc_name", ""),
        "POC_SUMMARY": poc.get("summary", ""),
        "GUIDE_FORKED": poc.get("guide_title", ""),
        "GUIDE_URL": poc.get("guide_url", ""),
        "FEATURES": poc.get("features") or [],
        "READINESS_SCORE": int(poc.get("readiness") or 0),
        "CONSIDERATIONS": poc.get("considerations") or [],
        "FIRST_STEP": poc.get("first_step", ""),
        "DOCUMENT_URL": state.get("blueprint_url", ""),
        "DELIVERY_STATUS": delivery_status(state),
        "DURATION_SECONDS": int(time.time() - started),
        "COCO_SECONDS": int(state.get("coco_seconds") or 0),
        "INPUT_TOKENS": int(state.get("input_tokens") or 0),
        "OUTPUT_TOKENS": int(state.get("output_tokens") or 0),
        "SE_OPERATOR": (cfg.get("event") or {}).get("operator", ""),
        "NOTES": f"{len(turns)} turns",
        # The qualification payload. Everything above is a pick from a list we
        # wrote; these two are the visitor's own words and their real estate.
        "PROBLEM_STATEMENT": vis.get("problem", ""),
        "PLATFORMS": state.get("platforms") or [],
        # Home-stage sovereignty answers - the visitor's own words on where they
        # are and what their rules are, useful for follow-up grouping.
        "COMPANY_COUNTRY": vis.get("company_country", ""),
        "RESIDENCY": vis.get("residency", ""),
        # QA-bot verdict for this blueprint. The per-finding detail is in
        # QA_FINDINGS; these four are the summary. QA_PASSED/QA_RELEVANT may be
        # NULL (bot off, or relevance undetermined), which the ::BOOLEAN cast of
        # a JSON null preserves.
        "QA_PASSED": qa.get("passed"),
        "QA_REPAIRS": int(qa.get("repairs") or 0),
        "QA_RELEVANT": qa.get("relevant"),
        "QA_NOTE": qa.get("note", ""),
    }
    bool_cols = ["QA_PASSED", "QA_RELEVANT"]
    arr_cols = ["DATA_HELD", "MARKETPLACE_JOINED", "FEATURES",
                "CONSIDERATIONS", "PLATFORMS"]
    cols = [k for k in payload if k not in arr_cols and k not in bool_cols]
    num = {"READINESS_SCORE", "DURATION_SECONDS", "COCO_SECONDS",
           "INPUT_TOKENS", "OUTPUT_TOKENS", "QA_REPAIRS"}
    sel = [f"p:{c}::{'NUMBER' if c in num else 'TEXT'}" for c in cols]
    sel += [f"p:{c}::ARRAY" for c in arr_cols]
    sel += [f"p:{c}::BOOLEAN" for c in bool_cols]

    sql = (f"INSERT INTO {tbl} (SESSION_TS, {', '.join(cols + arr_cols + bool_cols)}) "
           f"SELECT CURRENT_TIMESTAMP(), {', '.join(sel)} "
           f"FROM (SELECT PARSE_JSON(%s) AS p)")
    return sf_exec(cfg, sql, (json.dumps(payload, ensure_ascii=False),))


def log_qa(cfg, session_id, findings):
    """Write one QA_FINDINGS row per check that fired. This is the audit trail
    behind the silent repairs: the visitor's document is fixed in place, but the
    change is recorded here with its before/after. Best-effort - a logging
    failure must never block a visitor's handover. Returns (ok, error)."""
    if not findings:
        return True, ""
    sf = cfg.get("snowflake") or {}
    tbl = f"{sf.get('database')}.{sf.get('schema')}.QA_FINDINGS"
    rows, ok_all, last_err = 0, True, ""
    for f in findings:
        payload = {
            "SESSION_ID": session_id or "",
            "CHECK_NAME": str(f.get("check") or "")[:120],
            "SEVERITY": str(f.get("severity") or "")[:20],
            "DETAIL": str(f.get("detail") or "")[:2000],
            "REPAIRED": bool(f.get("repaired")),
            "BEFORE_VAL": f.get("before"),
            "AFTER_VAL": f.get("after"),
        }
        sql = (f"INSERT INTO {tbl} (SESSION_ID, FINDING_TS, CHECK_NAME, SEVERITY, "
               f"DETAIL, REPAIRED, BEFORE_VAL, AFTER_VAL) SELECT p:SESSION_ID::TEXT, "
               f"CURRENT_TIMESTAMP(), p:CHECK_NAME::TEXT, p:SEVERITY::TEXT, "
               f"p:DETAIL::TEXT, p:REPAIRED::BOOLEAN, p:BEFORE_VAL::TEXT, "
               f"p:AFTER_VAL::TEXT FROM (SELECT PARSE_JSON(%s) AS p)")
        ok, err = sf_exec(cfg, sql, (json.dumps(payload, ensure_ascii=False),))
        if ok:
            rows += 1
        else:
            ok_all, last_err = False, err
    return ok_all, last_err


def sf_exec(cfg, sql, args=None):
    """Run one statement on the shared connection with real bind parameters.

    Binds are why the base64 dance is gone: the driver ships values out of band,
    so braces, quotes and ampersands in a visitor's text cannot reach the SQL
    parser at all. It is also faster, since there is no `snow sql` subprocess.
    """
    try:
        cur = sf_conn(cfg).cursor()
        try:
            cur.execute(sql, args or ())
        finally:
            cur.close()
        return True, ""
    except Exception as e:                                       # noqa: BLE001
        return False, str(e)[:300].replace("\n", " ")


def delivery_status(state):
    """Truthful by construction. The QR to the presigned .docx is the delivery,
    so a staged URL is DELIVERED; anything else has not reached the visitor."""
    if state.get("blueprint_url"):
        return "DELIVERED"
    return "FAILED"


def log_turn(cfg, session_id, loc_id, text, reply, meta):
    """One row per location per visitor, so the slowest and most expensive beat
    is answerable rather than guessed at."""
    sf = cfg.get("snowflake") or {}
    tbl = f"{sf.get('database')}.{sf.get('schema')}.TURNS"
    u = (meta or {}).get("usage") or {}
    payload = {
        "SESSION_ID": session_id or "",
        "LOCATION": loc_id,
        "VISITOR_INPUT": (text or "")[:2000],
        "REPLY": (reply or "")[:2000],
        "DURATION_SECONDS": int((meta or {}).get("seconds") or 0),
        "INPUT_TOKENS": int(u.get("input_tokens") or 0),
        "OUTPUT_TOKENS": int(u.get("output_tokens") or 0),
        "CACHE_READ_TOKENS": int(u.get("cache_read_input_tokens") or 0),
        "SUCCEEDED": bool((meta or {}).get("ok")),
    }
    sql = (f"INSERT INTO {tbl} (TURN_TS, SESSION_ID, LOCATION, VISITOR_INPUT, "
           f"REPLY, DURATION_SECONDS, INPUT_TOKENS, OUTPUT_TOKENS, "
           f"CACHE_READ_TOKENS, SUCCEEDED) SELECT CURRENT_TIMESTAMP(), "
           f"p:SESSION_ID::TEXT, p:LOCATION::TEXT, p:VISITOR_INPUT::TEXT, "
           f"p:REPLY::TEXT, p:DURATION_SECONDS::NUMBER, p:INPUT_TOKENS::NUMBER, "
           f"p:OUTPUT_TOKENS::NUMBER, p:CACHE_READ_TOKENS::NUMBER, "
           f"p:SUCCEEDED::BOOLEAN FROM (SELECT PARSE_JSON(%s) AS p)")
    return sf_exec(cfg, sql, (json.dumps(payload, ensure_ascii=False),))


# ------------------------------------------------------------------- transports

class QRDelivery:
    """The only delivery path: build the .docx, stage it, presign it. The QR the
    visitor scans points at that presigned URL, which expires in seven days.

    There is deliberately NO local record. Email was removed, so the old outbox
    that wrote a fully-composed message (name, company, problem) to disk on every
    visit had no purpose left - it was pure liability, PII accumulating on a
    shared booth laptop between visitors. The governed Snowflake SESSIONS/TURNS
    row is the one intended persistence; nothing else about a visitor is kept.
    """

    last_meta = None

    def deliver(self, cfg, state, job_id):
        docx = blueprint_docx(cfg, state)
        url, err = ("", "")
        if docx:
            url, err = stage_and_presign(cfg, docx)
        if url:
            write_state({"blueprint_url": url})
        vis = state.get("visitor") or {}
        name = (vis.get("first_name") or "").strip()
        who = f", {name}" if name else ""
        if url:
            reply = (f"Wrapped and labelled{who}. Scan the code on screen and "
                     f"it is yours - the link works for seven days.")
            return True, reply
        # Presign failed: be honest and point at a person, do not claim delivery.
        return False, ("I could not wrap it up this time - grab a Snowflake "
                       "person and we will sort it." + (f" ({err})" if err else ""))


def get_transport(cfg):
    # One transport now. The setting is kept so a future hosted-delivery build
    # can register another, but the booth path is QR-to-presigned-stage only.
    return QRDelivery()


# ------------------------------------------------------------------ turn runner

def finish_turn(cfg, loc_id, label, reply, extra=None, meta=None):
    cur = read_state()
    turns = list(cur.get("turns") or [])
    turns.append({"location": loc_id, "text": label, "reply": reply,
                  "at": time.time()})
    patch = {"thinking": False, "reply": reply, "turns": turns}
    if meta:
        u = meta.get("usage") or {}
        patch["coco_seconds"] = int((cur.get("coco_seconds") or 0) + (meta.get("seconds") or 0))
        patch["input_tokens"] = int((cur.get("input_tokens") or 0) + (u.get("input_tokens") or 0))
        patch["output_tokens"] = int((cur.get("output_tokens") or 0) + (u.get("output_tokens") or 0))
    if extra:
        patch.update(extra)
    write_state(patch)
    # Always record the turn, even when no exec ran (e.g. the postbox no longer
    # spends an exec on a doomed draft attempt) - otherwise the location is
    # missing from TURNS and the per-beat timing table has a hole.
    ok, err = log_turn(cfg, cur.get("session_id"), loc_id, label, reply, meta)
    if not ok:
        print(f"[loco4coco] TURNS insert failed for {loc_id}: {err}")


def unlock_next(cfg, loc_id, state):
    order = cfg.get("unlock_order") or []
    if loc_id not in order:
        return state.get("unlocked") or []
    unlocked = list(state.get("unlocked") or [])
    i = order.index(loc_id)
    if i + 1 < len(order) and order[i + 1] not in unlocked:
        unlocked.append(order[i + 1])
    return unlocked


def run_checklist(cfg, loc_id, labels, job_id):
    state = read_state()
    key = "held" if loc_id == "library" else "joined"
    patch = {key: labels}
    if loc_id == "marketplace":
        # Resolve each pick back to its verified listing so the blueprint can
        # name the provider and link the listing. Anything typed into "Other"
        # stays a plain string - we will not guess a provider for it.
        ind = (state.get("visitor") or {}).get("industry") or "other"
        by_title = {r["title"]: r for r in listings_for(cfg, ind, state.get("session_id"), state)}
        patch["joined_listings"] = [by_title[l] for l in labels if l in by_title]
    write_state(patch)
    state = read_state()
    loc = (cfg.get("locations") or {}).get(loc_id) or {}
    body = fill(loc.get("prompt", ""), cfg, state,
                selection=", ".join(labels) or "nothing")
    prompt = "\n".join(base_context(cfg, state)) + "\n\n" + body
    ok, reply, meta = run_turn(cfg, prompt, loc_id, loc=loc, job_id=job_id)
    if not ok and not reply:
        return
    finish_turn(cfg, loc_id, ", ".join(labels), reply,
                {"unlocked": unlock_next(cfg, loc_id, state)}, meta)


def run_workshop(cfg, text, job_id):
    state = read_state()
    loc = (cfg.get("locations") or {}).get("workshop") or {}
    feature_list = ", ".join(sorted(load_features()))
    body = fill(loc.get("prompt", ""), cfg, state, input=text,
                feature_list=feature_list)
    # Deterministic shortlist appended to the prompt: the likely POC shape, the
    # listings for their industry, and the features that shape actually uses.
    #
    # This is how the closed lists reach `cortex exec` - as content, not as a
    # tool. exec takes no tools except through MCP, which is not guaranteed on a
    # borrowed booth laptop, and at ~220 tokens injecting the slice is cheaper
    # than any lookup would be. The full feature list above is deliberately kept
    # as well: narrowing it would risk excluding the right feature when the
    # archetype guess is wrong, and tokens are not the bottleneck here - the
    # ~18s process startup is.
    try:
        vis = (state.get("visitor") or {})
        hint_block = bctx.shortlist_block(
            " ".join([str(vis.get("problem") or ""), str(text or "")]),
            industry=vis.get("industry"), conn=coco_connection())
        if hint_block:
            body += ("\n\nFor reference, resolved server-side from their own "
                     "words:\n" + hint_block)
    except Exception:
        pass  # Guidance only. The prompt above already stands on its own.
    # Optional invitation to actually USE a tool. Measured: --bypass alone
    # produces NO tool calls, because the feature list is injected and guides
    # are resolved server-side, so CoCo has no reason to reach for one. Inviting
    # a docs lookup does fire a real tool but cost ~34s on its own, which would
    # undo the time the fast path just saved. Off by default; the Ask stop is
    # where tool use is genuinely warranted and budgeted for.
    hint = loc.get("tool_hint")
    if hint:
        body += "\n\n" + str(hint)
    prompt = "\n".join(base_context(cfg, state)) + "\n\n" + body
    ok, raw, meta = run_turn(cfg, prompt, "workshop", loc=loc, job_id=job_id)
    if not ok and not raw:
        return

    poc, reply = {}, raw
    m = re.search(r"\{.*\}", raw or "", re.S)
    if m:
        try:
            poc = json.loads(m.group(0))
        except json.JSONDecodeError:
            poc = {}
    if poc:
        reply = poc.get("reply") or "Right, I have what I need."
        arche = poc.get("archetype") or ""
        title, url = guide_for(cfg, arche)
        if not title:
            # Fall back on the most broadly useful fork rather than invent one.
            title, url = guide_for(cfg, "talk-to-my-data")
        poc["guide_title"], poc["guide_url"] = title, url
        # Drop anything outside the closed list, so every feature has a link.
        poc["features"] = [n for n, _u in link_features(poc.get("features"))]
        cons = poc.get("considerations")
        if isinstance(cons, str):
            cons = [cons]
        poc["considerations"] = [str(c).strip() for c in (cons or [])
                                if str(c).strip()][:4]
        try:
            poc["readiness"] = max(1, min(5, int(poc.get("readiness") or 3)))
        except (TypeError, ValueError):
            poc["readiness"] = 3
        # Stage 1 is deliberately thin - archetype, name, features, reply - which
        # is what took 84s and 1051 output tokens when it also had to write the
        # summary and four considerations. Fill the rest from the catalogue now
        # so the blueprint is already complete and correct, then improve it in
        # the background while the visitor walks to the postbox.
        d = archetype_defaults(cfg, arche)
        if not poc.get("features"):
            poc["features"] = d["features"]
        if not poc.get("first_step"):
            poc["first_step"] = d["first_step"]
        if not poc.get("summary"):
            poc["summary"] = ""
        if not poc.get("considerations"):
            poc["considerations"] = d["pool"][:3]
    else:
        poc = {"poc_name": text[:60], "summary": raw[:300], "features": [],
               "readiness": 3,
               "considerations": ["We ran out of time to work through the "
                                  "detail. Start by confirming what data you "
                                  "can actually get hold of."]}
        t, u = guide_for(cfg, "talk-to-my-data")
        poc["guide_title"], poc["guide_url"] = t, u

    finish_turn(cfg, "workshop", text, reply,
                {"poc": poc, "unlocked": unlock_next(cfg, "workshop", state),
                 "poc_pending": bool(poc.get("archetype"))}, meta)
    if poc.get("archetype"):
        threading.Thread(target=refine_poc, args=(cfg, text), daemon=True).start()


def refine_poc(cfg, text):
    """Stage 2. Runs AFTER CoCo has spoken, on the fast transport, while the
    visitor reads the reply and walks to the postbox. Fills the summary and the
    honest readiness score, and picks considerations from the precomputed pool by
    INDEX rather than writing them.

    Never blocks and never fails loudly: the blueprint is already complete from
    catalogue defaults, so a failure here costs tailoring, not the visit.
    """
    try:
        state = read_state()
        poc = dict(state.get("poc") or {})
        arche = poc.get("archetype") or ""
        d = archetype_defaults(cfg, arche)
        pool = d["pool"]
        if not pool:
            write_state({"poc_pending": False})
            return
        listed = "\n".join(f"{i + 1}. {c}" for i, c in enumerate(pool))
        vis = state.get("visitor") or {}
        prompt = (
            "You are helping shape a Snowflake proof of concept. Be concrete and "
            "never generic.\n\n"
            f"Person: {vis.get('first_name') or 'they'} at "
            f"{vis.get('company') or 'their organisation'}, "
            f"{industry_name(cfg, vis.get('industry'))}.\n"
            f"They hold: {', '.join(state.get('held') or []) or 'unspecified'}.\n"
            f"They will attach: {', '.join(state.get('joined') or []) or 'nothing yet'}.\n"
            f"They asked the POC to: \"{text}\"\n"
            f"It is a '{arche}' build called '{poc.get('poc_name') or 'their POC'}'.\n\n"
            "Candidate considerations:\n" + listed + "\n\n"
            "Return ONLY minified JSON, no prose, no code fence, with exactly "
            "these keys: {\"summary\": 2 sentences on what it does and why it "
            "matters to THEM, \"first_step\": one concrete first action, "
            "\"considerations\": array of exactly 3 integers, the numbers of the "
            "most relevant candidates above, \"readiness\": integer 1-5}\n"
            "Score readiness strictly: only award 4 or 5 if they named specific "
            "data, a specific question and a specific user. Nobody sees the "
            "number, so be honest rather than kind."
        )
        loc = {"transport": "complete"}
        ok, raw, _meta = run_turn(cfg, prompt, "workshop_refine", loc=loc)
        if not ok and not raw:
            write_state({"poc_pending": False})
            return
        m = re.search(r"\{.*\}", raw or "", re.S)
        extra = {}
        if m:
            try:
                extra = json.loads(m.group(0))
            except json.JSONDecodeError:
                extra = {}
        # Re-read: the visitor may have moved on and state may have changed.
        poc = dict((read_state().get("poc") or {}))
        if extra.get("summary"):
            poc["summary"] = str(extra["summary"])[:600]
        if extra.get("first_step"):
            poc["first_step"] = str(extra["first_step"])[:300]
        idx = extra.get("considerations")
        if isinstance(idx, list):
            picked = []
            for i in idx:
                try:
                    n = int(i)
                except (TypeError, ValueError):
                    continue
                if 1 <= n <= len(pool) and pool[n - 1] not in picked:
                    picked.append(pool[n - 1])
            if picked:
                poc["considerations"] = picked[:4]
        try:
            poc["readiness"] = max(1, min(5, int(extra.get("readiness")
                                                or poc.get("readiness") or 3)))
        except (TypeError, ValueError):
            pass
        write_state({"poc": poc, "poc_pending": False})
    except Exception:                                             # noqa: BLE001
        # A background refinement must never take the booth down.
        write_state({"poc_pending": False})


def run_ask(cfg, text, job_id):
    """The one unscripted moment: the visitor asks CoCo anything about their POC.

    This is where real CoCo genuinely earns it. Unlike the other turns there is
    no injected answer to reflect back, so CoCo has an actual reason to reach
    for its tools, and the tray shows real work. Always exec, never the fast
    path. Optional and skippable so it cannot blow the five minutes.
    """
    state = read_state()
    a = cfg.get("ask") or {}
    poc = state.get("poc") or {}
    body = fill(a.get("prompt", ""), cfg, state, input=text,
                poc_name=poc.get("poc_name") or "their POC",
                poc_summary=poc.get("summary") or "")
    hint = a.get("tool_hint")
    if hint:
        body += "\n\n" + str(hint)
    prompt = "\n".join(base_context(cfg, state)) + "\n\n" + body
    ok, reply, meta = run_turn(cfg, prompt, "ask", loc=a, job_id=job_id)
    if not ok and not reply:
        return
    finish_turn(cfg, "ask", text, reply, {"ask_used": True}, meta)


def qa_review(cfg, state, poc):
    """Review the finished blueprint against the request, repair what is safely
    fixable from the closed lists, and record every change.

    Returns (poc, findings, summary). Repairs are applied to a copy of `poc`;
    the visitor never sees a failure. Findings and the summary are written to
    Snowflake for audit (see run_send) - the "log that a change was made" rule.

    Deterministic checks run always; a model relevance turn runs when
    qa.model_review is on and fails open (never blocks the handover).
    """
    poc = dict(poc or {})
    vis = state.get("visitor") or {}
    findings = []

    def note(check, severity, detail, repaired=False, before=None, after=None):
        findings.append({"check": check, "severity": severity, "detail": detail,
                         "repaired": bool(repaired),
                         "before": None if before is None else str(before)[:200],
                         "after": None if after is None else str(after)[:200]})

    # 1. Every feature must resolve to a doc link, or it is dropped.
    before_feats = list(poc.get("features") or [])
    kept = [n for n, _u in link_features(before_feats)]
    if kept != before_feats:
        dropped = [f for f in before_feats if f not in kept]
        note("features_in_closed_list", "repair",
             "dropped feature(s) with no doc link: " + ", ".join(dropped),
             repaired=True, before=before_feats, after=kept)
        poc["features"] = kept

    # 2. The guide must resolve; if it fell back, say so.
    arche = poc.get("archetype") or ""
    title, url = guide_for(cfg, arche)
    if not title:
        title, url = guide_for(cfg, "talk-to-my-data")
        note("guide_resolves", "repair",
             "no fork for archetype %r; fell back to talk-to-my-data" % arche,
             repaired=True, before=arche, after="talk-to-my-data")
    if poc.get("guide_title") != title:
        poc["guide_title"], poc["guide_url"] = title, url

    # 3. Does the archetype match the visitor's own words? Flag, do not override -
    # the model saw the whole sentence and the resolver only scores keywords.
    problem = " ".join([str(vis.get("problem") or ""),
                        str(poc.get("_workshop_input") or "")]).strip()
    if arche and problem:
        try:
            hit, _ranked = bctx.resolve_archetype(problem, conn=coco_connection())
            if hit and hit.get("id") and hit["id"] != arche:
                note("archetype_matches_pain", "flag",
                     "model chose %r; pain text resolves to %r"
                     % (arche, hit["id"]))
        except Exception:
            pass

    # 4. Platform selection is normalised at ingest; re-apply the canonical rule
    # and flag if anything was still off - a cheap guard against a hand-edited
    # state or a future client sending raw values.
    plats = state.get("platforms") or []
    norm = normalise_platforms(cfg, plats)
    if list(norm) != list(plats):
        note("platforms_normalised", "flag",
             "platform set re-normalised for the blueprint",
             before=plats, after=norm)

    # 5. Readiness in range.
    try:
        r = int(poc.get("readiness") or 3)
    except (TypeError, ValueError):
        r = 3
    if r != poc.get("readiness"):
        note("readiness_in_range", "repair", "readiness coerced to 1-5",
             repaired=True, before=poc.get("readiness"), after=max(1, min(5, r)))
    poc["readiness"] = max(1, min(5, r))

    # 6. The visitor must leave with a name and a summary.
    if not str(poc.get("poc_name") or "").strip():
        note("poc_name_present", "flag", "no POC name produced")
    if not str(poc.get("summary") or "").strip():
        note("summary_present", "flag", "no summary produced")

    # 7. If they stated a residency rule, the sovereignty section must be there.
    residency = (vis.get("residency") or "").strip().lower()
    if residency in ("country_only", "eu") and not sovereignty_lines(cfg, state):
        note("sovereignty_consistent", "flag",
             "residency is %r but no sovereignty lines were produced" % residency)

    # 8. Model relevance turn - fail open.
    qacfg = cfg.get("qa") or {}
    relevant = None
    if qacfg.get("model_review") and problem:
        relevant = _qa_relevance(cfg, problem, poc)
        if relevant is False:
            note("relevance", "flag",
                 "model judged the POC does not address the stated problem")
        elif relevant is None:
            note("relevance_skipped", "info",
                 "relevance check did not complete; delivery proceeded")

    repairs = sum(1 for f in findings if f["repaired"])
    flags = sum(1 for f in findings if f["severity"] == "flag")
    summary = {"passed": flags == 0 and relevant is not False,
               "repairs": repairs, "relevant": relevant,
               "note": "; ".join(f["check"] for f in findings)[:400]}
    return poc, findings, summary


def _qa_relevance(cfg, problem, poc):
    """One fast COMPLETE call: does the POC address the stated problem? Returns
    True / False / None (None = could not tell, treated as non-blocking)."""
    qacfg = cfg.get("qa") or {}
    model = qacfg.get("model") or (cfg.get("coco") or {}).get("complete_model")
    prompt = (
        "You are a strict reviewer. A booth built this proof-of-concept plan for "
        "a visitor. Judge ONLY whether it addresses the problem they stated.\n\n"
        "Their problem: \"" + problem[:600] + "\"\n\n"
        "The POC: name=\"" + str(poc.get("poc_name") or "")[:120] + "\"; "
        "builds=\"" + str(poc.get("summary") or "")[:400] + "\"; "
        "features=" + ", ".join(poc.get("features") or []) + "\n\n"
        "Reply with ONLY minified JSON, nothing else: "
        '{"relevant": true or false, "issue": "at most 12 words, empty if relevant"}')
    loc = {"transport": "complete", "complete_model": model,
           "timeout": int(qacfg.get("timeout_seconds") or 12)}
    try:
        ok, raw, _meta = run_turn(cfg, prompt, "qa", loc=loc,
                                  timeout=loc["timeout"])
    except Exception:
        return None
    if not ok or not raw:
        return None
    m = re.search(r"\{.*\}", raw, re.S)
    if not m:
        return None
    try:
        v = json.loads(m.group(0))
    except json.JSONDecodeError:
        return None
    return bool(v.get("relevant")) if "relevant" in v else None


def run_send(cfg, job_id):
    state = read_state()

    # QA bot runs BEFORE delivery so its repairs land in the document the visitor
    # actually scans. It reviews the finished POC against their stated problem,
    # fixes what is safely fixable from the closed lists, and records every change.
    # Config-gated; if it errors it must never block the handover.
    findings = []
    if (cfg.get("qa") or {}).get("enabled", True):
        try:
            poc2, findings, summary = qa_review(cfg, state, state.get("poc") or {})
            write_state({"poc": poc2, "qa": summary})
            state = read_state()
        except Exception as e:                                   # noqa: BLE001
            print(f"[loco4coco] QA review error (non-blocking): {e}")

    transport = get_transport(cfg)
    delivered, reply = transport.deliver(cfg, state, job_id)
    st = read_state()
    write_state({"queued": bool(delivered), "email_sent": False})

    # Log BEFORE finish_turn. finish_turn clears `thinking`, which is the signal
    # everything else waits on, so logging after it means the row lands after
    # the visitor has already been told the flow is finished - and a failure
    # would go unnoticed until someone counted rows.
    ok, err = log_session(cfg, read_state())
    if ok:
        write_state({"logged": True, "log_error": ""})
    else:
        write_state({"logged": False, "log_error": err or "unknown"})
        print(f"[loco4coco] SESSIONS insert FAILED: {err}")

    # Audit trail for the silent repairs. Best-effort - a logging failure here
    # must not surface to the visitor, but it is printed so the operator sees it.
    qok, qerr = log_qa(cfg, read_state().get("session_id") or "", findings)
    if not qok:
        print(f"[loco4coco] QA_FINDINGS insert FAILED: {qerr}")

    finish_turn(cfg, "postbox", "send the blueprint", reply, {"stage": "done"},
                getattr(transport, "last_meta", None))


# ------------------------------------------------------------------ HTTP server

class Handler(BaseHTTPRequestHandler):
    server_version = "Loco4CoCo"
    # Timestamp of the last HTTP request. The idle watchdog uses it to stop a
    # forgotten server so it can never linger (and never burns anything: the
    # server is idle when no browser is polling).
    last_request = time.time()

    def log_message(self, fmt, *args):
        if "/api/state" not in (self.path or ""):
            super().log_message(fmt, *args)

    def _send(self, code, body=b"", ctype="text/plain; charset=utf-8"):
        Handler.last_request = time.time()
        self.send_response(code)
        self.send_header("Content-Type", ctype)
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        if body:
            self.wfile.write(body)

    def _json(self, obj, code=200):
        self._send(code, json.dumps(obj).encode("utf-8"), MIME[".json"])

    def _body(self):
        try:
            n = int(self.headers.get("Content-Length") or 0)
            return json.loads(self.rfile.read(n) or b"{}")
        except (ValueError, json.JSONDecodeError):
            return None

    def do_GET(self):
        path = urlparse(self.path).path
        if path == "/api/state":
            return self._json(read_state())
        if path == "/api/config":
            return self._json(load_config())
        if path == "/api/options":
            cfg, st = load_config(), read_state()
            return self._json({loc: options_for(cfg, loc, st)
                               for loc in (cfg.get("locations") or {})})
        if path == "/api/blueprint":
            # The visitor must leave with something even if the email never
            # arrives, so the finished blueprint is readable on screen.
            cfg, st = load_config(), read_state()
            poc = st.get("poc") or {}
            # readiness is internal lead-ranking data. This payload is rendered
            # straight to the visitor, so the score never leaves the server.
            safe_poc = {k: v for k, v in poc.items()
                        if k not in ("readiness", "weakest_point", "reply")}
            return self._json({
                "poc": safe_poc,
                "held": st.get("held") or [],
                "joined": st.get("joined") or [],
                "listings": st.get("joined_listings") or [],
                "features": [{"name": n, "url": u}
                             for n, u in link_features(poc.get("features"))],
                "considerations": poc.get("considerations") or [],
                "prompt": build_coco_prompt(cfg, st) if st.get("poc") else "",
                "document_url": st.get("blueprint_url") or "",
                "signup_url": (cfg.get("delivery") or {}).get("signup_url", ""),
                "email": (st.get("visitor") or {}).get("email", ""),
                "queued": bool(st.get("queued")),
                "draft_created": bool(st.get("draft_created")),
            })
        if path == "/api/qr":
            # Tier 2 of delivery: the presigned URL as a scannable SVG so the
            # visitor leaves with the document even if no email ever goes out.
            # Rendered server-side with segno (pure Python) rather than a
            # vendored JS encoder, so what they scan comes from a library we
            # can verify rather than hand-rolled bit placement.
            st = read_state()
            url = (parse_qs(urlparse(self.path).query).get("u") or
                   [st.get("blueprint_url") or ""])[0]
            if not url:
                return self._send(404, b"no document url yet")
            try:
                import io
                import segno
                buf = io.BytesIO()
                segno.make(url, error="M").save(
                    buf, kind="svg", scale=5, border=2,
                    dark="#0A121A", light="#FFFFFF")
                return self._send(200, buf.getvalue(), MIME[".svg"])
            except ImportError:
                return self._send(501, b"segno not installed")
            except Exception as e:                                # noqa: BLE001
                return self._send(500, str(e)[:200].encode("utf-8"))
        if path == "/blueprint":
            # The HTML blueprint, served by US with a real content type.
            # WHY THIS EXISTS: a presigned stage URL always serves
            # application/octet-stream (measured - GET_PRESIGNED_URL gives no way
            # to set it), so a QR pointing at the staged .html makes a phone
            # DOWNLOAD the file instead of rendering it. The QR therefore points
            # at the .docx, and the same markup renders here, on the booth screen,
            # and as the body of the email. This is the surface to replace the
            # day there is an SPCS or Streamlit endpoint to host it on.
            try:
                cfg, st = load_config(), read_state()
                page = blueprint_page(cfg, st)
                if not page:
                    return self._send(404, b"no blueprint yet")
                with open(page, "rb") as fh:
                    return self._send(200, fh.read(), MIME[".html"])
            except Exception as e:                                # noqa: BLE001
                return self._send(500, str(e)[:200].encode("utf-8"))
        if path == "/api/delivery/check":
            return self._delivery_check()
        if path in ("/", "/index.html"):
            return self._file("index.html")
        safe = posixpath.normpath(path).lstrip("/")
        if not safe or ".." in safe:
            return self._send(404, b"not found")
        return self._file(safe)

    def do_POST(self):
        path = urlparse(self.path).path
        routes = {
            "/api/intake": self._intake,
            "/api/home": self._home,
            "/api/select": self._select,
            "/api/compose": self._compose,
            "/api/ask": self._ask,
            "/api/send": self._post_it,
            "/api/reset": self._reset,
            "/api/state": self._patch,
            "/api/card": self._card,
        }
        fn = routes.get(path)
        if not fn:
            return self._send(404, b"not found")
        return fn()

    def _card(self):
        """Stage the share card and presign it, so it can leave on a phone.

        The card is drawn client-side on a canvas, and the browser's own
        `a.download` was the original handover - which cannot work: the booth
        machine is a shared demo laptop the visitor never logs into and walks
        away from (see CONSTRAINTS.md). The bytes therefore come back here as a
        data URL, get written as a real PNG, staged and presigned, and go home as
        a QR code. A presigned .png serves as octet-stream like everything else,
        which for an image on a phone is exactly right: it saves to Files or
        Photos and is theirs.
        """
        body = self._body() or {}
        raw = (body.get("png") or "").strip()
        if "," in raw:
            raw = raw.split(",", 1)[1]
        if not raw:
            return self._json({"error": "no image"}, 400)
        st = read_state()
        if st.get("card_url"):                     # idempotent, like POST IT
            return self._json({"url": st["card_url"]})
        try:
            png = base64.b64decode(raw, validate=True)
        except (ValueError, binascii.Error):
            return self._json({"error": "not base64"}, 400)
        if len(png) > 4_000_000 or png[:8] != b"\x89PNG\r\n\x1a\n":
            return self._json({"error": "not a png"}, 400)
        sid = st.get("session_id") or "visitor"
        name = "card-%s.png" % re.sub(r"[^A-Za-z0-9_-]", "", sid)[:40]
        path = os.path.join(tempfile.gettempdir(), name)
        with open(path, "wb") as fh:
            fh.write(png)
        url, err = stage_and_presign(load_config(), path)
        if not url:
            print("[loco] card presign failed:", err or "?")
            return self._json({"error": err or "presign failed"}, 502)
        write_state({"card_url": url})
        return self._json({"url": url})

    def _reset(self):
        # replace=True, so nothing from the previous visitor survives.
        st = write_state(dict(BLANK_STATE), replace=True)
        # Clear the visitor-named blueprint files from the OS temp dir too.
        purge_temp_artifacts()
        # Recycle the warm agent too. It probed stateless between calls, but a
        # shared booth must not merely trust that: restarting the process on
        # reset makes it impossible for one visitor's words to reach the next,
        # even if a future CLI build changed the agent's memory behaviour. Only
        # if it is actually running - do not spawn one when warming is off. Done
        # off-thread so reset stays instant; it re-warms on the next turn.
        try:
            ag = agent_pool.get_agent()
            if ag.alive():
                threading.Thread(target=ag.restart, daemon=True).start()
        except Exception:                                        # noqa: BLE001
            pass
        return self._json(st)

    def _patch(self):
        patch = self._body()
        if not isinstance(patch, dict):
            return self._json({"error": "expected an object"}, 400)
        return self._json(write_state(patch))

    def _intake(self):
        b = self._body()
        if not isinstance(b, dict):
            return self._json({"error": "invalid json"}, 400)
        cfg = load_config()
        first = (b.get("first_name") or "").strip()[:80]
        company = (b.get("company") or "").strip()[:120]
        email = (b.get("email") or "").strip()[:160]
        # Two sentences, generously bounded. This is the single highest-value
        # field on the form: it is the only place the visitor says WHY, and it
        # arrives before they have walked anywhere.
        problem = (b.get("problem") or "").strip()[:400]
        # Email is optional and no longer collected by the letter. Keep
        # accepting it so an operator-run intake can still supply one.
        if email and "@" not in email:
            email = ""
        if not first:
            return self._json({"error": "need a first name"}, 400)
        industry = b.get("industry") or infer_industry(cfg, company)
        if industry not in (cfg.get("industries") or {}):
            industry = "other"
        st = write_state({
            "visitor": {"first_name": first, "company": company,
                        "email": email, "industry": industry,
                        "problem": problem},
            "stage": "library", "unlocked": ["library"], "reasoning": [],
            "session_id": uuid.uuid4().hex,
        })
        # Fire Tier 0 now, while the visitor still has THE LIBRARY ahead of
        # them - see start_agentic_search()/listings_for(). Never delays
        # this response: it's a daemon thread, and a no-op if disabled.
        start_agentic_search(cfg, st.get("session_id"), industry, problem)
        return self._json({"ok": True, "industry": industry,
                           "industry_name": industry_name(cfg, industry),
                           "state": st})

    def _home(self):
        """Home-stage answers: where the data lives, where the company is based,
        and the residency rule. Captured before the Library now (moved off the
        Library's second question 2026-08-25). Enriches the blueprint; never
        gates correctness, so a failure here is not fatal to the visit.
        """
        b = self._body()
        if not isinstance(b, dict):
            return self._json({"error": "invalid json"}, 400)
        plats = [str(x)[:60] for x in (b.get("platforms") or []) if str(x).strip()]
        country = (b.get("company_country") or "").strip()[:60]
        # A closed vocabulary, so the blueprint's region logic can rely on it.
        residency = (b.get("residency") or "").strip().lower()
        if residency not in ("country_only", "eu", "us_ok", "unsure"):
            residency = "unsure"
        cur = read_state()
        vis = dict(cur.get("visitor") or {})
        vis["company_country"] = country
        vis["residency"] = residency
        st = write_state({"platforms": plats[:8], "visitor": vis})
        return self._json({"ok": True, "state": st})

    def _select(self):
        b = self._body()
        if not isinstance(b, dict):
            return self._json({"error": "invalid json"}, 400)
        cfg = load_config()
        loc_id = b.get("location") or ""
        if loc_id not in (cfg.get("locations") or {}):
            return self._json({"error": f"unknown location {loc_id!r}"}, 400)
        labels = [str(x)[:120] for x in (b.get("labels") or []) if str(x).strip()]
        # Platforms are captured on the home stage now (see _home), not here.
        if not labels:
            return self._json({"error": "nothing selected"}, 400)
        job = uuid.uuid4().hex
        loc = (cfg.get("locations") or {}).get(loc_id) or {}
        write_state({"thinking": True, "reply": "", "reasoning": [],
                     "transport": loc.get("transport") or "exec",
                     "job_id": job, "location": loc_id, "stage": loc_id})
        threading.Thread(target=run_checklist,
                         args=(cfg, loc_id, labels, job), daemon=True).start()
        return self._json({"accepted": True, "job_id": job})

    def _compose(self):
        b = self._body()
        if not isinstance(b, dict):
            return self._json({"error": "invalid json"}, 400)
        text = (b.get("text") or "").strip()[:500]
        if not text:
            return self._json({"error": "empty message"}, 400)
        cfg = load_config()
        job = uuid.uuid4().hex
        wloc = (cfg.get("locations") or {}).get("workshop") or {}
        write_state({"thinking": True, "reply": "", "reasoning": [],
                     "transport": wloc.get("transport") or "exec",
                     "job_id": job, "location": "workshop", "stage": "workshop"})
        threading.Thread(target=run_workshop, args=(cfg, text, job),
                         daemon=True).start()
        return self._json({"accepted": True, "job_id": job})

    def _ask(self):
        b = self._body()
        if not isinstance(b, dict):
            return self._json({"error": "invalid json"}, 400)
        text = (b.get("text") or "").strip()[:300]
        if not text:
            return self._json({"error": "empty question"}, 400)
        cfg = load_config()
        if not (cfg.get("ask") or {}).get("enabled", True):
            return self._json({"error": "ask is disabled"}, 400)
        st = read_state()
        if not st.get("poc"):
            return self._json({"error": "no POC to ask about yet"}, 400)
        job = uuid.uuid4().hex
        write_state({"thinking": True, "reply": "", "reasoning": [],
                     "transport": (cfg.get("ask") or {}).get("transport", "exec"),
                     "job_id": job, "location": "ask", "stage": "ask"})
        threading.Thread(target=run_ask, args=(cfg, text, job),
                         daemon=True).start()
        return self._json({"accepted": True, "job_id": job})

    def _post_it(self):
        cfg = load_config()
        st = read_state()
        # No email gate. The letter no longer asks for one: the QR handover is
        # the delivery that cannot fail, and the outbox draft is a bonus an
        # operator drains when there IS an address. Requiring one here is what
        # made SEND MY POC dead-end with a 400 on every press.
        if not st.get("poc"):
            return self._json({"error": "no POC to send yet"}, 400)
        # Already delivered for this visitor: hand back the same result rather
        # than re-staging the document and logging a duplicate SESSIONS row.
        if st.get("queued") and st.get("blueprint_url"):
            return self._json({"accepted": True, "already": True,
                               "blueprint_url": st.get("blueprint_url")})
        # Stage 2 may still be running if the visitor sprinted here. Give it a
        # moment, then go with what we have: the blueprint is already complete
        # from catalogue defaults, so this is a quality wait, not a hard one.
        waited = 0.0
        while read_state().get("poc_pending") and waited < 8.0:
            time.sleep(0.25)
            waited += 0.25
        job = uuid.uuid4().hex
        write_state({"thinking": True, "reply": "", "reasoning": [],
                     "transport": "none",
                     "job_id": job, "location": "postbox", "stage": "postbox"})
        threading.Thread(target=run_send, args=(cfg, job), daemon=True).start()
        return self._json({"accepted": True, "job_id": job})

    def _delivery_check(self):
        """Per-stand preflight. Names the leg that is broken rather than failing
        silently at the postbox with a visitor watching."""
        cfg = load_config()
        d = cfg.get("delivery") or {}
        checks = []

        def add(name, ok, detail=""):
            checks.append({"check": name, "ok": bool(ok), "detail": detail})

        try:
            import docx                                          # noqa: F401
            add("word_library", True, "python-docx importable")
        except Exception as e:                                    # noqa: BLE001
            add("word_library", False, f"pip install python-docx ({e})")

        stage = d.get("stage") or ""
        add("stage_configured", bool(stage), stage or "delivery.stage is empty")
        if stage:
            rows, err = snow_sql(cfg, f"LIST {stage}")
            add("stage_reachable", err == "", err or f"{len(rows or [])} objects")

        # End-to-end proof: stage a tiny file and presign it. This is the exact
        # path a visitor's document takes, so a green here means Tier 2 works.
        try:
            probe = os.path.join(tempfile.gettempdir(), "loco4coco-preflight.txt")
            with open(probe, "w", encoding="utf-8") as f:
                f.write("loco4coco preflight\n")
            url, err = stage_and_presign(cfg, probe)
            add("presign_works", bool(url), err or "presigned URL returned")
        except Exception as e:                                    # noqa: BLE001
            add("presign_works", False, str(e)[:160])

        add("delivery_transport", True,
            "QR to the presigned .docx is the handover; no email, no local record")

        ok = all(c["ok"] for c in checks
                 if c["check"] != "delivery_transport")
        return self._json({"ok": ok, "checks": checks,
                           "stage": stage,
                           "transport": d.get("transport")})

    def _file(self, rel):
        full = os.path.join(HERE, rel)
        if not os.path.isfile(full):
            return self._send(404, b"not found")
        ext = os.path.splitext(full)[1].lower()
        with open(full, "rb") as f:
            self._send(200, f.read(), MIME.get(ext, "application/octet-stream"))


def main():
    cfg = load_config()
    ap = argparse.ArgumentParser()
    ap.add_argument("--host", default=cfg["server"]["host"])
    ap.add_argument("--port", type=int, default=cfg["server"]["port"])
    args = ap.parse_args()

    if not os.path.exists(STATE_PATH):
        write_state(dict(BLANK_STATE))

    srv = ThreadingHTTPServer((args.host, args.port), Handler)
    print(f"Loco4CoCo companion running at http://{args.host}:{args.port}/")
    print(f"  event     : {cfg['event']['city']} ({cfg['event']['language']})")
    print(f"  account   : {cfg['snowflake']['connection_name']}")
    print(f"  transport : {(cfg.get('delivery') or {}).get('transport')}")
    print(f"  guides    : {len(load_guides())} primary forks loaded")
    # Stand identity is the one field that cannot be reconstructed after the
    # event. Every row is stamped with it at write time, so an empty value means
    # a day's conversations arrive with no way to tell which stand produced them.
    # Warn loudly rather than refuse: a booth that will not start is worse than
    # one with a gap in its reporting.
    _op = (cfg.get("event") or {}).get("operator", "")
    if str(_op).strip():
        print(f"  stand     : {_op}")
    else:
        print("  stand     : *** event.operator IS EMPTY *** every visitor row "
              "will be unattributable. Set it in game/config.json now - it "
              "cannot be added afterwards.")
    # Warm the model so the first real visitor does not pay cold-start latency.
    #
    # This used to spawn a throwaway `cortex exec`, which warmed almost nothing:
    # exec has no --resume or --session, so the next call was a cold process
    # again regardless. Now it starts the persistent `cortex mcp serve` process
    # and takes one cheap turn on it, because the first call on a fresh process
    # is measurably slower than the rest (~9.4s against ~6.0s, and ~3.4s once
    # settled). That cost belongs before doors open, not on visitor one.
    if (cfg.get("coco") or {}).get("warm_up", True):
        def _warm():
            c = cfg.get("coco") or {}
            if c.get("warm_agent", True):
                agent = agent_pool.get_agent(
                    connection=coco_connection(), model=c.get("model"),
                    workdir=HERE, log=lambda m: print("[loco4coco] " + m))
                ok, note = agent.warm()
                print("  warm-up   : warm agent %s (%s)"
                      % ("ready" if ok else "FAILED", note))
                if ok:
                    return
                # Fall through: if the pool will not start, the booth is about to
                # run on `cortex exec`, so warm whatever that can benefit from.
            cmd = [c.get("binary", "cortex"), "exec",
                   "Reply with the single word: ready.", "--no-mcp"]
            conn = (cfg.get("snowflake") or {}).get("connection_name")
            if conn:
                cmd += ["--connection", conn]
            if c.get("model"):
                cmd += ["-m", str(c["model"])]
            try:
                subprocess.run(cmd, stdout=subprocess.DEVNULL,
                               stderr=subprocess.DEVNULL, cwd=HERE, timeout=90)
            except Exception:                                    # noqa: BLE001
                pass
        threading.Thread(target=_warm, daemon=True).start()
        print("  warm-up   : dispatched")
    # Open the Snowflake connection now so the first visitor's fast-path turn
    # pays ~2s, not the ~5s that includes connection setup.
    def _warm_conn():
        try:
            sf_conn(cfg)
        except Exception:                                        # noqa: BLE001
            pass
    threading.Thread(target=_warm_conn, daemon=True).start()
    # Idle watchdog: if nothing hits the server for this many minutes (no browser
    # open, no test), it shuts itself down so it can never run for hours
    # unattended. Set server.idle_shutdown_minutes to 0 to disable.
    idle_min = (cfg.get("server") or {}).get("idle_shutdown_minutes", 45)
    if idle_min:
        def _watchdog():
            while True:
                time.sleep(30)
                if time.time() - Handler.last_request > idle_min * 60:
                    print(f"\nidle for {idle_min} min - shutting down.")
                    srv.shutdown()
                    return
        threading.Thread(target=_watchdog, daemon=True).start()
        print(f"  idle stop : after {idle_min} min with no requests")
    print("Ctrl-C to stop.")
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        print("\nstopped")


if __name__ == "__main__":
    main()
